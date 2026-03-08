"""
Company Pipeline & Trials Review Dashboard
Data: Company_Pipeline_Summary.csv (company), Biotech_Pipeline_Master.csv (asset), Enriched_Clinical_Trials.csv (trials)
Created by Tony Jiang
"""

import io
import json
import os
import pathlib
import re
from datetime import datetime
from urllib.parse import quote_plus

import pandas as pd
import streamlit as st
import yfinance as yf
from dotenv import load_dotenv
from google import genai

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None  # python-docx not installed

try:
    import docx2txt as _docx2txt
except ImportError:
    _docx2txt = None  # docx2txt not installed (handles legacy .doc via same API)

try:
    from striprtf.striprtf import rtf_to_text as _rtf_to_text
except ImportError:
    _rtf_to_text = None  # striprtf not installed

load_dotenv(pathlib.Path(__file__).parent / ".env")

# 优先读 Streamlit Secrets（Cloud 部署），回退到环境变量（本地 .env）
def _get_secret(key: str) -> str:
    try:
        return st.secrets.get(key, "") or os.environ.get(key, "")
    except Exception:
        return os.environ.get(key, "")

_GEMINI_API_KEY = _get_secret("GEMINI_API_KEY")
_gemini_client = genai.Client(api_key=_GEMINI_API_KEY) if _GEMINI_API_KEY else None
GEMINI_FLASH_MODEL = "gemini-2.5-flash"

# ── Page config & custom CSS ───────────────────────────────────

st.set_page_config(
    page_title="Biotech Intelligence Dashboard | Tony Jiang",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown(
    """
    <style>
    /* Dashboard title in header bar - white text on blue */
    header[data-testid="stHeader"] {
        background: linear-gradient(90deg, #1e3a5f 0%, #2d5a87 100%);
    }
    header[data-testid="stHeader"]::before {
        content: "Biotech Intelligence Dashboard for Non-oncology Companies";
        color: #ffffff;
        font-size: 1.35rem;
        font-weight: 600;
        display: block;
        padding: 0.5rem 0;
        white-space: nowrap;
    }
    /* Maximize main content area — padding-top 留出 header 高度，确保 Tab 栏可见 */
    .block-container { padding-top: 3.5rem; padding-bottom: 0.5rem; padding-left: 1.5rem; padding-right: 1.5rem; max-width: 100%; }
    h1 { font-size: 1.75rem !important; white-space: nowrap; }
    .stDataFrame { font-size: 0.9rem; }
    /* Compact metrics */
    [data-testid="stMetricValue"] { font-size: 1.1rem; }
    /* Footer */
    .footer { text-align: right; font-size: 0.8rem; color: #6b7280; margin-top: 1rem; padding-top: 0.5rem; border-top: 1px solid #e5e7eb; }
    /* Section headers */
    .section-title { font-size: 1.05rem; font-weight: 600; color: #1e3a5f; margin-bottom: 0.25rem; }
    .company-detail-header { background: #f0f4f8; padding: 0.6rem 1rem; border-radius: 6px; margin-bottom: 0.5rem; border-left: 4px solid #2d5a87; }
    /* 侧边栏展开/收起按钮：固定在左上角，橙色高亮确保任何背景下都清晰可见 */
    [data-testid="collapsedControl"] {
        background-color: #e07b00 !important;
        border-radius: 0 8px 8px 0 !important;
        opacity: 1 !important;
        z-index: 999 !important;
    }
    [data-testid="collapsedControl"]:hover {
        background-color: #f59e0b !important;
    }
    [data-testid="collapsedControl"] svg {
        fill: #ffffff !important;
        color: #ffffff !important;
    }
    /* 一键复制按钮样式 */
    .copy-btn-wrapper button {
        background-color: #2d5a87;
        color: white;
        border: none;
        border-radius: 6px;
        padding: 0.3rem 0.9rem;
        cursor: pointer;
        font-size: 0.85rem;
    }
    .copy-btn-wrapper button:hover {
        background-color: #1e3a5f;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

DATA_DIR = pathlib.Path(__file__).parent
SUMMARY_CSV = DATA_DIR / "Company_Pipeline_Summary.csv"
MASTER_CSV = DATA_DIR / "Biotech_Pipeline_Master.csv"
TRIALS_CSV = DATA_DIR / "Enriched_Clinical_Trials.csv"
ASSOCIATED_FILES_DIR = DATA_DIR / "associated_files"
ASSOCIATED_FILES_INDEX_CSV = DATA_DIR / "associated_files_index.csv"
AI_DD_REPORT_DIR = DATA_DIR / "AI_DD_REPORT"
UPSIDE_CACHE_FILE = DATA_DIR / "upside_cache.json"
PT_DIR = DATA_DIR / "Paper_Trading_Portfolios"
PT_HISTORY_CSV = PT_DIR / "portfolio_history.csv"
PT_ARENA_REPORTS_DIR = PT_DIR / "Arena_Reports"  # AI 解析用研报存档
PT_DIR.mkdir(exist_ok=True)
PT_ARENA_REPORTS_DIR.mkdir(exist_ok=True)


def _load_upside_cache() -> tuple[dict, str]:
    """从磁盘读取 Upside 缓存。返回 (upside_map, 上次更新时间字符串)。"""
    if UPSIDE_CACHE_FILE.exists():
        try:
            data = json.loads(UPSIDE_CACHE_FILE.read_text(encoding="utf-8"))
            return data.get("upside_map", {}), data.get("fetched_at", "")
        except Exception:
            pass
    return {}, ""


def _save_upside_cache(upside_map: dict, fetched_at: str) -> None:
    """将 Upside 数据持久化到磁盘。"""
    try:
        UPSIDE_CACHE_FILE.write_text(
            json.dumps({"upside_map": upside_map, "fetched_at": fetched_at}, ensure_ascii=False),
            encoding="utf-8",
        )
    except Exception:
        pass


def _ensure_company_files_dir(symbol: str) -> pathlib.Path:
    """Ensure associated_files / {Symbol} exists; return that path."""
    p = ASSOCIATED_FILES_DIR / symbol.strip().upper()
    p.mkdir(parents=True, exist_ok=True)
    return p


def _load_associated_index() -> pd.DataFrame:
    if not ASSOCIATED_FILES_INDEX_CSV.exists():
        df = pd.DataFrame(columns=["Symbol", "FilePath", "DisplayName", "UploadedAt"])
    else:
        df = pd.read_csv(ASSOCIATED_FILES_INDEX_CSV, encoding="utf-8-sig")
    # 每次加载时根据磁盘目录扫描，把已有文件并入 index（刷新后仍能认出）
    df, changed = _sync_associated_index_from_disk(df)
    if changed:
        _save_associated_index(df)
    return df


def _sync_associated_index_from_disk(existing: pd.DataFrame) -> tuple[pd.DataFrame, bool]:
    """扫描 associated_files/{Symbol}/ 下所有文件，将不在 index 中的加入。返回 (合并后的 df, 是否有新增)。"""
    if not ASSOCIATED_FILES_DIR.exists():
        return existing, False
    existing_paths: set[tuple[str, str]] = set()
    if not existing.empty and "Symbol" in existing.columns and "FilePath" in existing.columns:
        existing_paths = {
            (str(sym).strip().upper(), str(fp).strip().replace("\\", "/"))
            for sym, fp in zip(existing["Symbol"], existing["FilePath"])
        }
    new_rows = []
    for symbol_dir in ASSOCIATED_FILES_DIR.iterdir():
        if not symbol_dir.is_dir():
            continue
        sym = symbol_dir.name.upper()
        for f in symbol_dir.iterdir():
            if not f.is_file():
                continue
            rel = f"{sym}/{f.name}".replace("\\", "/")
            if (sym, rel) in existing_paths:
                continue
            try:
                uploaded_at = datetime.fromtimestamp(f.stat().st_mtime).isoformat()
            except Exception:
                uploaded_at = datetime.now().isoformat()
            new_rows.append({
                "Symbol": sym,
                "FilePath": rel,
                "DisplayName": f.name,
                "UploadedAt": uploaded_at,
            })
            existing_paths.add((sym, rel))
    if not new_rows:
        return existing, False
    return pd.concat([existing, pd.DataFrame(new_rows)], ignore_index=True), True


def _save_associated_index(df: pd.DataFrame) -> None:
    df.to_csv(ASSOCIATED_FILES_INDEX_CSV, encoding="utf-8-sig", index=False)


def _safe_filename(name: str) -> str:
    return re.sub(r"[^\w\-_.]", "_", name)[:80]


def _ensure_ai_dd_report_dir(symbol: str) -> pathlib.Path:
    """Ensure AI_DD_REPORT/{Symbol} exists; return that path."""
    p = AI_DD_REPORT_DIR / symbol.strip().upper()
    p.mkdir(parents=True, exist_ok=True)
    return p


def _build_dd_prompt(
    selected_symbol: str,
    company_name: str,
    market_cap: float,
    total_cash: float,
    pipeline_rows: pd.DataFrame,
) -> str:
    """组装终极 DD Prompt 模板。"""
    def _fmt(val, default="N/A"):
        if val is None or (isinstance(val, float) and pd.isna(val)):
            return default
        return str(val).strip() or default
    market_cap_str = f"${market_cap / 1e9:.2f}B" if isinstance(market_cap, (int, float)) and market_cap and not pd.isna(market_cap) else "N/A"
    total_cash_str = f"${total_cash / 1e6:.0f}M" if isinstance(total_cash, (int, float)) and total_cash and not pd.isna(total_cash) else "N/A"
    lines = []
    for _, row in pipeline_rows.iterrows():
        asset = _fmt(row.get("Asset_Name"))
        phase = _fmt(row.get("Highest_Phase"))
        moa = _fmt(row.get("Mechanism_of_Action"))
        catalyst = _fmt(row.get("Next_Catalyst_Date"))
        lines.append(f"- 药物/代号: {asset} | 阶段: {phase} | 机制: {moa} | 催化剂: {catalyst}")
    pipeline_str = "\n".join(lines) if lines else "(无管线数据)"
    return f"""你现在是一位横跨华尔街顶级投行（熟悉估值模型与并购逻辑）与顶尖药企临床研发（精通试验设计与统计学底牌）的资深 Biotech 量化投研专家。

你的任务是基于我提供的公司宏观基本面、临床管线数据、ATTACH 的材料，以及你自己通过 Deep Research 搜索到的全网最新情报（务必深度挖掘 SEC 10-K/10-Q 财报、财报电话会 Earnings Call 逐字稿、医学会议如 JPM/ASCO/ESMO/ASH 的最新 Data Readout），为我撰写一份极其硬核、冷酷、客观的投资尽职调查报告（DD Report）。

【输入数据上下文】
公司名称与代码：{selected_symbol} — {company_name}
当前市值：{market_cap_str}
账上现金：{total_cash_str}
核心管线列表：
{pipeline_str}

请严格按照以下 7 个模块输出研报，拒绝废话，直击投资痛点：

### 1. 管理层基本盘 (Management & Track Record)
- **团队背景**：核心高管与研发负责人的过往履历。过去是否有成功将药物推进至获批（FDA Approval）或成功将公司卖给 Big Pharma 的 Track Record？

### 2. 科学机制与临床/统计学推演 (Science & Clinical Data)
针对进度最快或价值最高的 1-2 个核心管线进行深度拆解：
- **MoA 护城河**：其作用机制（MoA）是 First-in-class，还是已有验证的 Validated target，抑或是完全 Novel 的盲盒？
- **试验设计审视**：拆解其核心临床试验。对照组（Comparator/SOC）选择是否合理？终点设计（Endpoints）是否存在水分？
- **数据透视**：挖掘其 Pre-clinical 与已公布的 Clinical 数据并进行严格对比。

### 3. 竞争格局与"中国 Biotech"防线 (Competitive Space & Moat)
- **赛道全景**：在同一适应症下，对比 In-class（同靶点）和 Out-of-class（治同种病）的竞品进度与数据。
- **反内卷护城河**：针对中国 Biotech 极速研发的 "Me-too/Fast-follower" 威胁，该公司的 Asset 是否具备足够深的护城河？

### 4. 商业化前景与市场空间 (Commercial Outlook & Unmet Need)
- **TAM 与痛点**：目标适应症的市场规模与真实的 Unmet Need。
- **商业化破局**：目前的 SOC 卖得如何？该公司准备如何实现差异化（Differentiate）？

### 5. 并购博弈与 Big Pharma 胃口 (M&A / BD Appetite)
- **历史先例**：梳理近期是否有 Big Pharma 收购类似靶点或适应症药物的先例与溢价。
- **潜在买家预测**：结合 Big Pharma 目前的管线短板，预测谁是最有可能的收购者？预估 Buyout Price 是多少？

### 6. 财务生存期与华尔街共识 (Financials, Valuation & Sentiment)
- **股价复盘与情绪**：分析过去一年股价异动原因，挖掘市场 Sentiment。
- **资金生死线**：基于最新 Burn Rate 和账上现金计算其实际 Runway。这笔钱能否撑到下一个数据读出之后？是否存在极高的数据公布前定增砸盘风险？
- **估值与目标价**：汇总华尔街 Target Price。结合 rNPV 或 DCF，给出一个大模型独立预测的 Target Price 并给出推导理由。

### 7. 最终审判局与情景假设 (The Final Verdict & Scenarios)
- **生死局总结**：判定估值是被严重低估还是高估？
- **情景推演**：牛市假设 (Bull Case) 能涨到多少？熊市假设 (Bear Case) 底线在哪？"""


def _list_report_files(symbol: str) -> list[tuple[str, str]]:
    """返回 [(文件名, 人类可读时间), ...]，按时间倒序。
    时间优先从文件名时间戳解析，失败则用文件修改时间（mtime）。
    """
    dir_path = AI_DD_REPORT_DIR / symbol.strip().upper()
    if not dir_path.exists() or not dir_path.is_dir():
        return []
    result = []
    for f in dir_path.iterdir():
        if not (f.is_file() and f.name.startswith("Report_") and f.suffix.lower() in (".md", ".txt", ".gdoc", ".doc", ".docx", ".rtf")):
            continue
        stem = f.stem
        ts_str = stem[7:] if stem.startswith("Report_") else ""
        human = None
        try:
            dt = datetime.strptime(ts_str, "%Y%m%d_%H%M%S")
            human = dt.strftime("%Y-%m-%d %H:%M:%S")
        except (ValueError, Exception):
            pass
        if human is None:
            # 回退：用文件系统修改时间
            try:
                human = datetime.fromtimestamp(f.stat().st_mtime).strftime("%Y-%m-%d %H:%M:%S") + "（文件时间）"
            except Exception:
                human = f.name
        result.append((f.name, human))
    result.sort(key=lambda x: x[0], reverse=True)
    return result


def _cell_to_md(s: str) -> str:
    """单元格内容转 Markdown 安全：换行改空格，管道符转全角避免破坏表格。"""
    if not s:
        return ""
    t = s.strip().replace("\n", " ").replace("\r", " ")
    return t.replace("|", "｜")  # 全角管道符，避免破坏 Markdown 表格


def _table_to_markdown(table) -> str:
    """将 python-docx Table 转为 Markdown 表格字符串。"""
    rows = []
    for row in table.rows:
        cells = [_cell_to_md(cell.text) for cell in row.cells]
        if any(cells):
            rows.append("| " + " | ".join(cells) + " |")
    if not rows:
        return ""
    # 表头分隔行
    n_cols = len(table.rows[0].cells)
    sep = "| " + " | ".join(["---"] * n_cols) + " |"
    return rows[0] + "\n" + sep + "\n" + "\n".join(rows[1:])


def _extract_docx_text(file_path: pathlib.Path) -> tuple[str, str | None]:
    """从 .docx 文件提取正文与表格，表格以 Markdown 形式保留。返回 (正文, 错误信息)。"""
    if DocxDocument is None:
        return "", "未安装 python-docx。请在终端运行: pip install python-docx"
    try:
        doc = DocxDocument(io.BytesIO(file_path.read_bytes()))
        parts = []
        # 按文档顺序：段落与表格交错（python-docx 的 paragraphs/tables 是分开的，需用 element.body 保序）
        from docx.text.paragraph import Paragraph
        from docx.table import Table
        for el in doc.element.body:
            tag = el.tag.split("}")[-1] if "}" in el.tag else el.tag
            if tag == "p":
                p = Paragraph(el, doc)
                if p.text and p.text.strip():
                    parts.append(p.text.strip())
            elif tag == "tbl":
                t = Table(el, doc)
                md = _table_to_markdown(t)
                if md:
                    parts.append(md)
        # 若上面保序失败（兼容旧版），则回退为段落+表格
        if not parts:
            for p in doc.paragraphs:
                if p.text and p.text.strip():
                    parts.append(p.text.strip())
            for table in doc.tables:
                md = _table_to_markdown(table)
                if md:
                    parts.append(md)
        # 页眉、页脚
        for section in doc.sections:
            for para in section.header.paragraphs:
                if para.text and para.text.strip():
                    parts.append(para.text.strip())
            for para in section.footer.paragraphs:
                if para.text and para.text.strip():
                    parts.append(para.text.strip())
        text = "\n\n".join(parts).strip()
        return text, None
    except Exception as e:
        return "", str(e)


def _extract_any_text(file_path: pathlib.Path) -> tuple[str, str | None]:
    """
    通用文件文本提取，支持：
      .md / .txt          — 直接读取
      .docx               — python-docx（含表格转 Markdown）
      .doc                — docx2txt（纯 Python，无需 Word）
      .rtf                — striprtf
      .gdoc               — Google Docs 快捷方式，提取 URL
    返回 (text, error_msg)，成功时 error_msg=None。
    """
    suf = file_path.suffix.lower()
    if suf in (".md", ".txt"):
        try:
            return file_path.read_text(encoding="utf-8", errors="replace").strip(), None
        except Exception as e:
            return "", str(e)
    if suf == ".docx":
        return _extract_docx_text(file_path)
    if suf == ".doc":
        if _docx2txt is None:
            return "", "未安装 docx2txt，请运行: pip install docx2txt"
        try:
            text = _docx2txt.process(str(file_path))
            return (text or "").strip(), None
        except Exception as e:
            return "", f".doc 解析失败：{e}"
    if suf == ".rtf":
        if _rtf_to_text is None:
            return "", "未安装 striprtf，请运行: pip install striprtf"
        try:
            raw = file_path.read_bytes().decode("utf-8", errors="replace")
            text = _rtf_to_text(raw)
            return (text or "").strip(), None
        except Exception as e:
            return "", f".rtf 解析失败：{e}"
    if suf == ".gdoc":
        try:
            data = json.loads(file_path.read_text(encoding="utf-8", errors="replace"))
            url = data.get("url", "") if isinstance(data, dict) else ""
            return f"[Google Doc 链接]({url})" if url else "(Google Doc，无法提取 URL)", None
        except Exception as e:
            return "", f".gdoc 解析失败：{e}"
    return "", f"不支持的格式：{suf}"


def _gemini_summarize_report(report_text: str) -> tuple[str | None, str | None]:
    """调用 Gemini Flash 对研报做总结。返回 (总结文本, 错误信息)，成功时错误为 None。"""
    if not _gemini_client:
        return None, "未配置 GEMINI_API_KEY，请在 .env 中设置。"
    if not (report_text or "").strip():
        return None, "当前研报无正文可总结（如为 .doc 或 .gdoc 请先下载查看后复制内容）。"
    prompt = """你是一位 Biotech 投研专家。请对以下尽职调查报告（DD Report）做简明总结。

要求：
- 用中文输出
- 保留核心结论、主要风险与投资要点
- 分点或分段落，条理清晰
- 控制在 500 字以内

报告原文：
---
"""
    prompt += (report_text or "").strip()
    try:
        resp = _gemini_client.models.generate_content(
            model=GEMINI_FLASH_MODEL,
            contents=prompt,
        )
        text = (resp.text or "").strip()
        return (text, None) if text else (None, "API 返回为空")
    except Exception as e:
        return None, str(e)


# Same control-token filter as build_pipeline (for display-only cleaning)
_CTRL_PATTERN = re.compile(
    r"\b(placebo|sham|standard\s+of\s+care|soc|best\s+supportive\s+care|bsc|"
    r"saline|normal\s+saline|inactive|comparator|control|dummy)\b",
    re.IGNORECASE,
)
# 明确跳过 “MATCHING PLACEBO / PLACEBO MATCHING” 等变体（不限长度）
_PLACEBO_VARIANT_PATTERN = re.compile(
    r"\bmatching\s+placebo\b|\bplacebo\s+matching\b|\bplacebo\s+for\s+|\bplacebo\s+of\s+|\bplacebo\s+to\s+match|"
    r"\bplacebo\s+control\b|\bplacebo\s+comparator\b|\bthe\s+placebo\b|"
    r"\bmatched\s+placebo\b|\bplacebo\s+matched\b|\bplacebo\s+in\s+",
    re.IGNORECASE,
)
# 用于展示时过滤 Master 中的 placebo 类 asset（与 build_pipeline 的 is_placebo_asset 逻辑一致）
_PLACEBO_ASSET_PATTERN = re.compile(
    r"^(.*\bmatching\s+placebo\b|.*\bplacebo\s+matching\b|"
    r"placebo\s+for\s+|placebo\s+to\s+match|placebo\s+of\s+|placebo\s*\(|"
    r".*\bplacebo\s+control\b|.*\bplacebo\s+comparator\b|.*\bthe\s+placebo\b|"
    r".*\bplacebo\s+matched\b|.*\bmatched\s+placebo\b|.*\bplacebo\s+in\s+|"
    r".*\bplacebo\s+(booster|vaccine|pen\b|injector|syringe|DPI|auto-injector)\b|"
    r".*vehicle\s*\(?\s*placebo|.*\s+or\s+placebo\s*$|.*/placebo\s*$|"
    r".*-placebo\s*$|.*\s+placebo\s*$|.*\bplacebo\b.*)$",
    re.IGNORECASE,
)
# Strip trailing dose specs (e.g. "0.5 mg", "10 mg/kg QD") so "ALKS 2680 0.5mg" -> "ALKS 2680"
_DOSE_SUFFIX = re.compile(
    r"\s+\d+(?:\.\d+)?\s*(?:mg|mcg|µg|g|kg|iu|ui|ml|%|mg/kg|mcg/kg|µg/kg|mg/day|mcg/day)(?:\s*(?:qd|bid|tid|daily|weekly|twice|once))?\s*$",
    re.IGNORECASE,
)
# Token that is only a dose (e.g. "4mg", "6 mg", "10mg QD") — skip these, we only want drug names
_DOSE_ONLY_PATTERN = re.compile(
    r"^\s*\d+(?:\.\d+)?\s*(?:mg|mcg|µg|g|kg|iu|ui|ml|%|mg/kg|mcg/kg|µg/kg|mg/day|mcg/day)(?:\s*(?:qd|bid|tid|daily|weekly|twice|once))?\s*$",
    re.IGNORECASE,
)
# 清洗 TREATMENT 类型前缀/后缀：OPEN LABEL, BLINDED, Open-Label, Double-Blind, High/Low Dose 等
_TREATMENT_STRIP_LEADING = re.compile(
    r"^\s*(?:open\s*[-]?\s*label\s+(?:extension\s*[-]?\s*period\s*)?|blinded\s+|double\s*[-]?\s*blind(?:ed)?\s+|single\s*[-]?\s*blind(?:ed)?\s+|"
    r"high\s+dose\s+of\s+|low\s+dose\s+of\s+|high\s+dose\s+|low\s+dose\s+)+",
    re.IGNORECASE,
)
_TREATMENT_STRIP_TRAILING = re.compile(
    r"\s+(?:open\s*[-]?\s*label|blinded|double\s*[-]?\s*blind|single\s*[-]?\s*blind|high\s+dose|low\s+dose)\s*$",
    re.IGNORECASE,
)


def _is_placebo_asset(name: str) -> bool:
    """与 build_pipeline 一致：若为 placebo/control 类 asset 则返回 True，展示时过滤。"""
    if not name or not str(name).strip():
        return False
    return bool(_PLACEBO_ASSET_PATTERN.match(str(name).strip()))


def _strip_treatment_prefix_suffix(s: str) -> str:
    """Strip OPEN LABEL / BLINDED / Open-Label / High Dose / Low Dose 等前缀与后缀，便于归一化显示。"""
    if not s or not str(s).strip():
        return s
    t = str(s).strip()
    while _TREATMENT_STRIP_LEADING.search(t):
        t = _TREATMENT_STRIP_LEADING.sub("", t, count=1).strip()
    while _TREATMENT_STRIP_TRAILING.search(t):
        t = _TREATMENT_STRIP_TRAILING.sub("", t, count=1).strip()
    return t.strip() or s


def _normalize_asset_for_grouping(asset_name: str) -> str:
    """
    与 build_pipeline 一致：同一药物的多种写法归一为同一 key（Blinded ESK-001 / Open-Label ESK-001 -> ESK-001），
    用于展示时按药物合并行。
    """
    if not asset_name or not str(asset_name).strip():
        return str(asset_name).strip() or ""
    t = _strip_treatment_prefix_suffix(str(asset_name).strip())
    t = re.sub(r"\s+", " ", t).strip()
    t = re.sub(r"([A-Za-z0-9])\s+(\d)", r"\1-\2", t)
    return t.strip() or str(asset_name).strip()


def _normalize_intervention_base(s: str) -> str:
    """Remove trailing dose suffix and treatment prefix/suffix so same drug maps to one name."""
    t = s.strip()
    while _DOSE_SUFFIX.search(t):
        t = _DOSE_SUFFIX.sub("", t).strip()
    t = _strip_treatment_prefix_suffix(t)
    return t


def _interventions_display(raw: str) -> str:
    """Strip placebo/control and dose-only tokens; show only drug names, deduplicated (no doses)."""
    if not raw or str(raw).strip() in ("", "N/A"):
        return ""
    parts = re.split(r"\s*[,|]\s*", str(raw))
    kept_bases = []
    seen_bases = set()
    for p in parts:
        p = p.strip()
        if not p:
            continue
        # 明确跳过 “MATCHING PLACEBO / PLACEBO MATCHING” 等变体（不限长度）
        if _PLACEBO_VARIANT_PATTERN.search(p):
            continue
        if _CTRL_PATTERN.search(p) and len(p) < 50:
            continue
        # Skip tokens that are purely a dose (e.g. "4mg", "6 mg", "10mg QD")
        if _DOSE_ONLY_PATTERN.match(p):
            continue
        base = _normalize_intervention_base(p)
        if not base:
            continue
        key = base.upper()
        if key in seen_bases:
            continue
        seen_bases.add(key)
        kept_bases.append(base)
    return ", ".join(kept_bases) if kept_bases else str(raw).strip()


@st.cache_data(show_spinner="Loading company summary…")
def load_summary() -> pd.DataFrame:
    if not SUMMARY_CSV.exists():
        return pd.DataFrame()
    df = pd.read_csv(SUMMARY_CSV)
    for col in ["Market Cap", "EV", "Total Debt", "Total Cash", "Price", "52W Low", "52W High", "Wall Street Ratings",
                "Pipeline_Count", "Total_Active_Trials", "Shares Outstanding", "Institutional Shares", "Insider %"]:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors="coerce")
    return df


@st.cache_data(show_spinner="Loading trials…")
def load_trials() -> pd.DataFrame:
    if not TRIALS_CSV.exists():
        return pd.DataFrame()
    df = pd.read_csv(TRIALS_CSV)
    for col in ["EnrollmentCount"]:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors="coerce")
    return df


@st.cache_data(show_spinner="Loading pipeline master…")
def load_master() -> pd.DataFrame:
    """中间层：公司 -> 药物管线 (Biotech_Pipeline_Master.csv)，不修改内容与结构，仅读取。"""
    if not MASTER_CSV.exists():
        return pd.DataFrame()
    return pd.read_csv(MASTER_CSV)


def get_realtime_financials(symbol: str) -> dict:
    """
    从 yfinance 获取单只股票的实时财务与华尔街共识数据。
    使用 .get() 安全提取，缺失填 'N/A'。返回 dict 含：
    Price, Market Cap, Total Cash, targetHighPrice, targetMeanPrice, targetLowPrice,
    recommendationKey（首字母大写美化）, numberOfAnalystOpinions。
    """
    if not symbol or str(symbol).strip() == "":
        return _empty_financials()
    sym = str(symbol).strip().upper()
    try:
        ticker = yf.Ticker(sym)
        info = ticker.info or {}
    except Exception:
        return _empty_financials()
    def _num(val, default="N/A"):
        if val is None:
            return default
        try:
            return float(val)
        except (TypeError, ValueError):
            return default
    # 评级：recommendationKey 如 "buy", "strong_buy" -> "Buy", "Strong Buy"
    rec_raw = info.get("recommendationKey") or info.get("recommendation")
    if rec_raw is not None and str(rec_raw).strip():
        rec = str(rec_raw).replace("_", " ").strip().title()
    else:
        rec = "N/A"
    return {
        "Price": _num(info.get("currentPrice") or info.get("regularMarketPrice")),
        "Market Cap": _num(info.get("marketCap")),
        "Total Cash": _num(info.get("totalCash")),
        "targetHighPrice": _num(info.get("targetHighPrice")),
        "targetMeanPrice": _num(info.get("targetMeanPrice")),
        "targetLowPrice": _num(info.get("targetLowPrice")),
        "recommendationKey": rec,
        "numberOfAnalystOpinions": _num(info.get("numberOfAnalystOpinions")),
    }


def _empty_financials() -> dict:
    return {
        "Price": "N/A",
        "Market Cap": "N/A",
        "Total Cash": "N/A",
        "targetHighPrice": "N/A",
        "targetMeanPrice": "N/A",
        "targetLowPrice": "N/A",
        "recommendationKey": "N/A",
        "numberOfAnalystOpinions": "N/A",
    }


def get_analyst_upside_batch(symbols: tuple) -> dict:
    """
    批量获取各 symbol 的 Upside %（(targetMeanPrice - currentPrice) / currentPrice * 100）。
    返回 dict: symbol -> float 或 None（无覆盖时为 None）。symbols 为 tuple 以便可哈希缓存。
    """
    result = {}
    for sym in symbols:
        if not sym or str(sym).strip() == "":
            result[sym] = None
            continue
        sym = str(sym).strip().upper()
        try:
            ticker = yf.Ticker(sym)
            info = ticker.info or {}
            price = info.get("currentPrice") or info.get("regularMarketPrice")
            target = info.get("targetMeanPrice")
            if price is not None and target is not None:
                try:
                    p, t = float(price), float(target)
                    if p > 0:
                        result[sym] = round((t - p) / p * 100.0, 2)
                    else:
                        result[sym] = None
                except (TypeError, ValueError):
                    result[sym] = None
            else:
                result[sym] = None
        except Exception:
            result[sym] = None
    return result


# ── Load data ─────────────────────────────────────────────────────

df_summary = load_summary()
df_trials_all = load_trials()
df_master = load_master()

if df_summary.empty:
    st.error(f"**{SUMMARY_CSV.name}** not found. Run the pipeline first.")
    st.stop()

# ── Sidebar: filters ──────────────────────────────────────────────

st.sidebar.markdown("### Filters")
st.sidebar.markdown("---")

# Market Cap (B)
mcap_b = df_summary["Market Cap"].dropna() / 1e9
mcap_min_b = float(mcap_b.min()) if len(mcap_b) else 0.0
mcap_max_b = float(mcap_b.max()) if len(mcap_b) else 1.0
st.sidebar.markdown("**Market Cap (B USD)**")
mcap_lo = st.sidebar.number_input("Min", min_value=0.0, max_value=mcap_max_b * 1.2, value=mcap_min_b, step=0.1, format="%.2f", key="mcap_lo")
mcap_hi = st.sidebar.number_input("Max", min_value=0.0, max_value=mcap_max_b * 2, value=mcap_max_b, step=0.5, format="%.2f", key="mcap_hi")

# Therapeutic Area
tas = sorted(df_summary["Therapeutic_Area_Filter"].dropna().unique().tolist())
TA_FILTER_ALLOWED = frozenset([
    "Cardiovascular", "Dermatology", "Gastroenterology", "Hematology",
    "Immunology", "Infectious Diseases", "Metabolic/Endocrinology", "Musculoskeletal",
    "Neurology/CNS", "No Trials", "Ophthalmology", "Others", "Pain", "Rare Disease",
    "Respiratory", "Urology/Nephrology",
])
tas = [t for t in tas if t in TA_FILTER_ALLOWED]
ta_selected = st.sidebar.multiselect("Therapeutic Area", options=tas, default=[], key="ta_filter")

# Phase
_phase_order = {"PHASE4": 4, "PHASE3": 3, "PHASE2": 2, "PHASE1": 1, "EARLY_PHASE1": 0, "N/A": -1, "No Trials": -2}
phases_raw = df_summary["Highest_Phase"].dropna().unique().tolist()
phases = sorted(phases_raw, key=lambda x: (_phase_order.get(x, -1), str(x)))
phase_selected = st.sidebar.multiselect("Phase of Development", options=phases, default=[], key="phase_filter")

# Has Marketed Drug
marketed_options = ["All", "Yes", "No"]
marketed_choice = st.sidebar.radio("Has Marketed Drug", options=marketed_options, index=0, key="marketed")

# Has DD Report
st.sidebar.markdown("**Has DD Report in System**")
report_filter_choice = st.sidebar.radio("Has DD Report", options=["All", "Yes", "No"], index=0, key="report_filter", label_visibility="collapsed")

# Upside % (华尔街目标价潜在涨幅，允许负数)
st.sidebar.markdown("**Upside %** (Target vs Current)")
upside_lo = st.sidebar.number_input("Min %", value=-999.0, step=1.0, format="%.1f", key="upside_lo", help="e.g. 10 = 10%+ upside; -999 = no lower bound")
upside_hi = st.sidebar.number_input("Max %", value=999.0, step=1.0, format="%.1f", key="upside_hi", help="e.g. 50 = up to 50%; 999 = no upper bound")

st.sidebar.markdown("---")
execute_btn = st.sidebar.button("Execute Search", type="primary", width="stretch")
clear_btn = st.sidebar.button("Clear filters", width="stretch")
refresh_btn = st.sidebar.button("🔄 Refresh data (reload CSV)", width="stretch", help="清缓存并重新读取 CSV。修改过 CSV 或跑过 normalize_ta_csvs 后请点此以看到最新 TA/公司数据。")
upside_refresh_btn = st.sidebar.button("📈 Refresh Upside % (yfinance)", width="stretch", help="手动从 yfinance 拉取最新华尔街目标价与 Upside %，结果存至磁盘，重启后仍可用。")
# 从磁盘读取上次拉取时间（刷新页面后仍显示）
_disk_upside_map, _disk_upside_fetched_at = _load_upside_cache()
if _disk_upside_fetched_at:
    st.sidebar.caption(f"Upside 上次更新：{_disk_upside_fetched_at}")
else:
    st.sidebar.caption("Upside %：尚未拉取，点上方按钮刷新")
st.sidebar.markdown("---")
st.sidebar.caption("Dashboard by **Tony Jiang**")

# Session state for applied filters (apply only on Execute Search)
if "search_executed" not in st.session_state:
    st.session_state["search_executed"] = False
if execute_btn:
    st.session_state["applied_mcap_lo"] = mcap_lo
    st.session_state["applied_mcap_hi"] = mcap_hi
    st.session_state["applied_ta"] = ta_selected
    st.session_state["applied_phase"] = phase_selected
    st.session_state["applied_marketed"] = marketed_choice
    st.session_state["applied_report_filter"] = report_filter_choice
    st.session_state["applied_upside_lo"] = upside_lo
    st.session_state["applied_upside_hi"] = upside_hi
    st.session_state["search_executed"] = True
    st.rerun()
if clear_btn:
    st.session_state["search_executed"] = False
    # 不能直接给 widget 的 key 赋值，否则会报错；删除 key 后 rerun，widget 会用默认值
    for key in ("mcap_lo", "mcap_hi", "ta_filter", "phase_filter", "marketed", "upside_lo", "upside_hi", "company_select"):
        if key in st.session_state:
            del st.session_state[key]
    st.rerun()
if refresh_btn:
    st.cache_data.clear()
    st.rerun()
if upside_refresh_btn:
    st.session_state["upside_fetch_requested"] = True
    st.rerun()


# Use applied filter values when Execute Search was clicked; otherwise default = show all companies
if st.session_state["search_executed"] and "applied_mcap_lo" in st.session_state:
    _mcap_lo = st.session_state["applied_mcap_lo"]
    _mcap_hi = st.session_state["applied_mcap_hi"]
    _ta = st.session_state["applied_ta"]
    _phase = st.session_state["applied_phase"]
    _marketed = st.session_state["applied_marketed"]
    _upside_lo = st.session_state.get("applied_upside_lo", -999.0)
    _upside_hi = st.session_state.get("applied_upside_hi", 999.0)
    _report_filter = st.session_state.get("applied_report_filter", "All")
else:
    # Default: show all companies (full range, no TA/phase/marketed/upside filter)
    _mcap_lo = mcap_min_b
    _mcap_hi = mcap_max_b
    _ta = []
    _phase = []
    _marketed = "All"
    _upside_lo = -999.0
    _upside_hi = 999.0
    _report_filter = "All"

# ── Apply filters to summary (default = all; after Execute Search = applied values) ─

filtered = df_summary.copy()
filtered = filtered[
    (filtered["Market Cap"].fillna(0) >= _mcap_lo * 1e9) &
    (filtered["Market Cap"].fillna(0) <= _mcap_hi * 1e9)
]
if _ta:
    filtered = filtered[filtered["Therapeutic_Area_Filter"].isin(_ta)]
if _phase:
    filtered = filtered[filtered["Highest_Phase"].isin(_phase)]
if _marketed == "Yes":
    filtered = filtered[filtered["Has_Marketed_Drug"] == "Yes"]
elif _marketed == "No":
    filtered = filtered[filtered["Has_Marketed_Drug"] == "No"]
# 研报过滤（需在 _report_date_map 建立之前先用临时扫描，此处直接扫目录）
if _report_filter != "All":
    _syms_with_report = set(_report_date_map.keys()) if "_report_date_map" in dir() else set()
    # _report_date_map 在 tabs 定义之前才建，此处先做一次快速扫描
    if not AI_DD_REPORT_DIR.exists():
        _syms_with_report = set()
    else:
        _syms_with_report = {
            d.name.upper() for d in AI_DD_REPORT_DIR.iterdir()
            if d.is_dir() and any(
                f.name.startswith("Report_") for f in d.iterdir() if f.is_file()
            )
        }
    if _report_filter == "Yes":
        filtered = filtered[filtered["Symbol"].str.upper().isin(_syms_with_report)]
    elif _report_filter == "No":
        filtered = filtered[~filtered["Symbol"].str.upper().isin(_syms_with_report)]
filtered = filtered.drop_duplicates(subset="Symbol", keep="first").reset_index(drop=True)

# Upside：点按钮时拉取并写入磁盘；每次启动从磁盘读取，刷新页面后仍保留上次数据
_upside_filter_active = _upside_lo > -999 or _upside_hi < 999

if st.session_state.get("upside_fetch_requested"):
    st.session_state["upside_fetch_requested"] = False
    symbols_tuple = tuple(filtered["Symbol"].dropna().astype(str).str.strip().str.upper().unique().tolist())
    if symbols_tuple:
        with st.spinner("正在拉取 yfinance Upside 数据，请稍候…"):
            _fetched_map = get_analyst_upside_batch(symbols_tuple)
        _fetched_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        _save_upside_cache(_fetched_map, _fetched_at)
        # 同步写入磁盘后重新读取，确保下方逻辑用最新数据
        _disk_upside_map, _disk_upside_fetched_at = _fetched_map, _fetched_at

# 使用磁盘缓存的 upside_map（刷新页面后仍有数据）
if _disk_upside_map:
    filtered["Upside_Pct"] = filtered["Symbol"].astype(str).str.strip().str.upper().map(_disk_upside_map)
else:
    filtered["Upside_Pct"] = float("nan")

# 仅当用户明确设置了非默认区间时才按 Upside 过滤
if _upside_filter_active:
    filtered = filtered[
        filtered["Upside_Pct"].notna() &
        (filtered["Upside_Pct"] >= _upside_lo) &
        (filtered["Upside_Pct"] <= _upside_hi)
    ].reset_index(drop=True)

# ── 扫描 AI_DD_REPORT，建立每家公司最新研报日期映射 ──────────────
def _build_report_date_map() -> dict[str, str]:
    """返回 {SYMBOL: 'YYYY-MM-DD'} 最新 Report_ 文件日期，无研报则不含该 key。"""
    result = {}
    if not AI_DD_REPORT_DIR.exists():
        return result
    for sym_dir in AI_DD_REPORT_DIR.iterdir():
        if not sym_dir.is_dir():
            continue
        sym = sym_dir.name.upper()
        latest_dt = None
        for f in sym_dir.iterdir():
            if not (f.is_file() and f.name.startswith("Report_") and
                    f.suffix.lower() in (".md", ".txt", ".gdoc", ".doc", ".docx", ".rtf")):
                continue
            ts_str = f.stem[7:]  # strip "Report_"
            dt = None
            try:
                dt = datetime.strptime(ts_str, "%Y%m%d_%H%M%S")
            except ValueError:
                try:
                    dt = datetime.fromtimestamp(f.stat().st_mtime)
                except Exception:
                    pass
            if dt and (latest_dt is None or dt > latest_dt):
                latest_dt = dt
        if latest_dt:
            result[sym] = latest_dt.strftime("%Y-%m-%d")
    return result

_report_date_map = _build_report_date_map()

# Unique companies for selector (from filtered rows)
companies = filtered[["Symbol", "Name"]].drop_duplicates().sort_values("Name")
company_list = ["— Select company —"] + companies.apply(lambda r: f"{r['Symbol']} — {r['Name']}", axis=1).tolist()

# ── 三大主页：st.tabs 导航 ────────────────────────────────────────
_tab_screener, _tab_portfolio, _tab_arena = st.tabs([
    "📊 公司筛选与研报",
    "💼 投资组合优化",
    "📈 策略竞技场",
])

with _tab_screener:
    st.caption("Use sidebar filters and **Execute Search** to apply · Select a company to filter list and view trials")
    st.markdown("---")

    # ── Top bar: KPI + Select Company ──────────────────────────────────

    col_k1, col_k2, col_k3, col_sel = st.columns([1, 1, 1, 2])
    with col_k1:
        st.metric("Companies", f"{len(filtered)}")
    with col_k2:
        avg_mcap = filtered["Market Cap"].mean() / 1e9
        st.metric("Avg Market Cap", f"${avg_mcap:.2f}B" if pd.notna(avg_mcap) else "—")
    with col_k3:
        with_catalyst = (filtered["Next_Catalyst"].notna()) & (filtered["Next_Catalyst"] != "") & (filtered["Next_Catalyst"] != "Passed")
        st.metric("With future catalyst", f"{with_catalyst.sum()}")
    with col_sel:
        company_choice = st.selectbox(
            "Select company",
            options=company_list,
            index=0,
            key="company_select",
        )

    # ── Display table: one row per company (or single company if selected) ────────
    display = filtered.copy()
    # When a company is selected, filter table to that company only
    if company_choice and company_choice != "— Select company —":
        sel_symbol = company_choice.split(" — ")[0]
        display = display[display["Symbol"] == sel_symbol].copy()

    # Format for display (keep originals for sorting)
    display["Market_Cap_B"] = (display["Market Cap"] / 1e9).round(2)
    display["EV_B"] = (display["EV"] / 1e9).round(2)
    display["Total_Debt_M"] = (display["Total Debt"] / 1e6).round(0)
    display["Total_Cash_M"] = (display["Total Cash"] / 1e6).round(0)
    # 研报日期列
    display["Latest Report"] = display["Symbol"].str.upper().map(_report_date_map)

    # Column order — Upside_Pct 紧跟 Price 方便排序；Latest Report 紧跟 Symbol
    cols_show = [
        "Symbol", "Latest Report", "Price", "Upside_Pct", "Market_Cap_B", "Name", "Therapeutic_Areas", "Has_Marketed_Drug",
        "Highest_Phase", "Pipeline_Count", "Total_Active_Trials", "Next_Catalyst",
        "Country", "52W Low", "52W High", "EV_B",
        "Shares Outstanding", "Institutional Shares", "Total_Debt_M", "Total_Cash_M",
        "Wall Street Ratings",
    ]
    cols_show = [c for c in cols_show if c in display.columns]
    display = display[cols_show]

    # Rename for display labels
    display = display.rename(columns={
        "Market_Cap_B": "Market Cap (B)",
        "EV_B": "EV (B)",
        "Total_Debt_M": "Total Debt (M)",
        "Total_Cash_M": "Total Cash (M)",
        "Upside_Pct": "Upside %",
    })

    # Table height: small when 1 company selected so Trial info comes up
    table_height = 100 if len(display) <= 1 else 520

    st.markdown(
        "<p class='section-title'>Company list (one row per company)"
        + (" — filtered to selected company" if company_choice and company_choice != "— Select company —" else "")
        + " — sort by clicking column headers</p>",
        unsafe_allow_html=True,
    )
    st.dataframe(
        display.reset_index(drop=True),
        width="stretch",
        hide_index=True,
        height=table_height,
        column_config={
            "Latest Report": st.column_config.TextColumn("📄 Latest Report", width="small"),
            "Price": st.column_config.NumberColumn("Price", format="$%.2f"),
            "Market Cap (B)": st.column_config.NumberColumn("Market Cap ($B)", format="%.2f"),
            "EV (B)": st.column_config.NumberColumn("EV ($B)", format="%.2f"),
            "52W Low": st.column_config.NumberColumn("52W Low", format="$%.2f"),
            "52W High": st.column_config.NumberColumn("52W High", format="$%.2f"),
            "Total Debt (M)": st.column_config.NumberColumn("Total Debt ($M)", format="%.0f"),
            "Total Cash (M)": st.column_config.NumberColumn("Total Cash ($M)", format="%.0f"),
            "Wall Street Ratings": st.column_config.NumberColumn("WS Rating", format="%.2f"),
            "Upside %": st.column_config.NumberColumn("Upside %", format="%.1f"),
            "Therapeutic_Areas": st.column_config.TextColumn("Therapeutic Areas", width="large"),
        },
    )

    # ── Company detail: 三级标签页（管线概览 | 研报生成器 | 研报阅读）───────────────────

    if company_choice and company_choice != "— Select company —":
        selected_symbol = company_choice.split(" — ")[0]
        company_name = company_choice.split(" — ", 1)[1] if " — " in company_choice else ""

        st.markdown("---")
        _sel_report_date = _report_date_map.get(selected_symbol.upper())
        _report_badge = f"&nbsp;·&nbsp; 📄 最新研报：{_sel_report_date}" if _sel_report_date else "&nbsp;·&nbsp; 📄 暂无研报"
        st.markdown(
            f"<div class='company-detail-header'><strong>{selected_symbol}</strong> — {company_name}{_report_badge}</div>",
            unsafe_allow_html=True,
        )

        tab1, tab2, tab3 = st.tabs(["📊 管线与临床概览", "🧠 智能研报生成器", "📂 研报阅读与投研笔记"])

        # ── Tab 1: 管线与临床概览 ─────────────────────────────────────────────
        with tab1:
            _company_row = None
            if selected_symbol and not df_summary[df_summary["Symbol"] == selected_symbol].empty:
                _company_row = df_summary[df_summary["Symbol"] == selected_symbol].iloc[0]
            _website = str(_company_row.get("Website", "")).strip() if _company_row is not None and "Website" in df_summary.columns else ""
            _ir_link = str(_company_row.get("IR_Search_Link", "")).strip() if _company_row is not None and "IR_Search_Link" in df_summary.columns else ""
            if _website in ("", "nan", "N/A"):
                _website = f"https://www.google.com/search?q={quote_plus(company_name)}"
            if _ir_link in ("", "nan", "N/A"):
                _ir_link = f"https://www.google.com/search?q={quote_plus(company_name + ' investor presentation PDF')}"
            _link_cols = st.columns([1, 1, 4])
            with _link_cols[0]:
                st.link_button("🌐 Official Website", _website)
            with _link_cols[1]:
                st.link_button("📄 IR / Investor Decks", _ir_link)

            with st.spinner("Loading analyst data…"):
                _fin = get_realtime_financials(selected_symbol)
            with st.expander("🏦 华尔街共识 (Wall Street Consensus)", expanded=True):
                _price = _fin.get("Price")
                _target_mean = _fin.get("targetMeanPrice")
                _has_target = isinstance(_target_mean, (int, float)) and _target_mean not in (None, "")
                _has_price = isinstance(_price, (int, float)) and _price not in (None, "") and (_price or 0) > 0
                if not _has_target or not _has_price:
                    st.info("No Analyst Coverage — 该公司暂无华尔街目标价/共识数据。")
                else:
                    _upside_pct = (_target_mean - _price) / _price * 100.0
                    c1, c2, c3, c4, c5, c6 = st.columns(6)
                    with c1:
                        st.metric("Current Price", f"${_price:.2f}" if isinstance(_price, (int, float)) else str(_price))
                    with c2:
                        _th = _fin.get("targetHighPrice")
                        st.metric("Target High", f"${_th:.2f}" if isinstance(_th, (int, float)) else str(_th))
                    with c3:
                        st.metric("Target Mean", f"${_target_mean:.2f}")
                    with c4:
                        _tl = _fin.get("targetLowPrice")
                        st.metric("Target Low", f"${_tl:.2f}" if isinstance(_tl, (int, float)) else str(_tl))
                    with c5:
                        st.metric("Consensus", _fin.get("recommendationKey", "N/A"))
                    with c6:
                        _n = _fin.get("numberOfAnalystOpinions")
                        st.metric("Analysts", str(int(_n)) if isinstance(_n, (int, float)) else str(_n))
                    _upside_str = f"{_upside_pct:+.1f}%"
                    if _upside_pct > 0:
                        st.markdown(f"<span style='color: #16a34a; font-weight: 600;'>Upside: {_upside_str}</span>", unsafe_allow_html=True)
                    else:
                        st.markdown(f"<span style='color: #b91c1c; font-weight: 600;'>Upside: {_upside_str}</span>", unsafe_allow_html=True)

            st.markdown("**Associated Files**")
            _company_dir = _ensure_company_files_dir(selected_symbol)
            _af_index = _load_associated_index()
            _files_for_symbol = _af_index[_af_index["Symbol"].astype(str).str.strip().str.upper() == selected_symbol.upper()] if not _af_index.empty and selected_symbol else pd.DataFrame()
            _up_col1, _up_col2 = st.columns([2, 1])
            with _up_col1:
                _uploaded = st.file_uploader("Upload any file to associate with this company", type=None, key=f"assoc_file_{selected_symbol}")
            with _up_col2:
                if _uploaded is not None and selected_symbol:
                    if st.button("Upload & associate with " + selected_symbol, type="primary", key=f"assoc_btn_{selected_symbol}"):
                        _ts = datetime.now().strftime("%Y%m%d_%H%M%S")
                        _stored_name = f"{_ts}_{_safe_filename(_uploaded.name)}"
                        _rel_path = f"{selected_symbol.upper()}/{_stored_name}"
                        _path = ASSOCIATED_FILES_DIR / selected_symbol.upper() / _stored_name
                        _path.write_bytes(_uploaded.getvalue())
                        _new = pd.DataFrame([{"Symbol": selected_symbol.upper(), "FilePath": _rel_path, "DisplayName": _uploaded.name, "UploadedAt": datetime.now().isoformat()}])
                        _af_index = _load_associated_index()
                        _save_associated_index(pd.concat([_af_index, _new], ignore_index=True))
                        st.success("Saved and associated with " + selected_symbol)
                        st.rerun()
            if not _files_for_symbol.empty:
                st.caption("Associated file(s) — download or delete:")
                for _idx, _row in _files_for_symbol.iterrows():
                    _rel = _row["FilePath"]
                    _fpath = ASSOCIATED_FILES_DIR / _rel
                    _dname = _row.get("DisplayName", _rel.split("/")[-1])
                    # 显示时间：优先 index 中 UploadedAt 字段，其次 mtime
                    _uploaded_at = str(_row.get("UploadedAt", "")).strip()
                    if not _uploaded_at or _uploaded_at == "nan":
                        try:
                            _uploaded_at = datetime.fromtimestamp(_fpath.stat().st_mtime).strftime("%Y-%m-%d %H:%M") if _fpath.exists() else ""
                        except Exception:
                            _uploaded_at = ""
                    _time_label = f"  🕐 {_uploaded_at}" if _uploaded_at else ""
                    _c1, _c2 = st.columns([3, 1])
                    with _c1:
                        if _fpath.exists():
                            _bytes = _fpath.read_bytes()
                            st.download_button(
                                label=f"📎 {_dname}{_time_label}",
                                data=_bytes, mime="application/octet-stream",
                                file_name=_dname, key=f"dl_af_{selected_symbol}_{_idx}",
                            )
                        else:
                            st.caption(f"Missing: {_rel}")
                    with _c2:
                        if st.button("Delete", key=f"del_af_{selected_symbol}_{_idx}"):
                            if _fpath.exists():
                                _fpath.unlink()
                            _af_index = _load_associated_index()
                            _save_associated_index(_af_index[_af_index["FilePath"] != _rel].reset_index(drop=True))
                            st.rerun()

            assets = pd.DataFrame()
            if not df_master.empty and "Symbol" in df_master.columns:
                assets = df_master.loc[df_master["Symbol"].astype(str).str.strip().str.upper() == selected_symbol.upper()].copy()
            _phase_order = {"PHASE4": 4, "PHASE3": 3, "PHASE2": 2, "PHASE1": 1, "EARLY_PHASE1": 0, "N/A": -1, "No Trials": -2}
            if not assets.empty and "Highest_Phase" in assets.columns:
                assets["_phase_ord"] = assets["Highest_Phase"].map(lambda x: _phase_order.get(x, -1))
                assets = assets.sort_values("_phase_ord", ascending=False).drop(columns=["_phase_ord"], errors="ignore")
            if not assets.empty and "Asset_Name" in assets.columns:
                assets = assets[~assets["Asset_Name"].apply(lambda x: _is_placebo_asset(x) if pd.notna(x) else False)]
            if not assets.empty:
                assets["_norm"] = assets["Asset_Name"].apply(lambda x: _normalize_asset_for_grouping(str(x)) if pd.notna(x) else "")
                assets = assets[assets["_norm"].astype(str).str.len() > 0].copy()
            if assets.empty:
                st.markdown("<p class='section-title'>药物管线 · Drugs</p>", unsafe_allow_html=True)
                st.info("未在 Biotech_Pipeline_Master.csv 中找到该公司的药物管线记录。")
            else:
                _phase_rank = {"PHASE4": 5, "PHASE3": 4, "PHASE2": 3, "PHASE1": 2, "EARLY_PHASE1": 1, "N/A": 0, "No Trials": -1}
                def _merge_assets(grp):
                    row = grp.iloc[0].to_dict()
                    row["Asset_Name"] = grp["_norm"].iloc[0]
                    if grp["Highest_Phase"].notna().any():
                        best = grp["Highest_Phase"].map(lambda p: _phase_rank.get(p, -1)).idxmax()
                        row["Highest_Phase"] = grp.loc[best, "Highest_Phase"]
                    row["Active_Trial_Count"] = grp["Active_Trial_Count"].apply(lambda x: int(x) if pd.notna(x) else 0).sum()
                    nct_list = []
                    for s in grp["Trial_NCTIds"].dropna():
                        for t in re.split(r"[,;]\s*", str(s)):
                            t = t.strip()
                            if t and (t.upper().startswith("NCT") or t.isdigit()):
                                nct_list.append(t if t.upper().startswith("NCT") else f"NCT{t}")
                    row["Trial_NCTIds"] = ", ".join(dict.fromkeys(nct_list))
                    cond_set = set()
                    for c in grp["Detailed_Conditions"].dropna():
                        for item in str(c).split(";"):
                            if item.strip():
                                cond_set.add(item.strip())
                    row["Detailed_Conditions"] = "; ".join(sorted(cond_set)) if cond_set else (grp["Detailed_Conditions"].iloc[0] if grp["Detailed_Conditions"].notna().any() else "")
                    dates = grp["Next_Catalyst_Date"].dropna().astype(str).str.strip()
                    dates = dates[~dates.isin(("", "N/A", "Passed"))]
                    row["Next_Catalyst_Date"] = dates.iloc[0] if len(dates) else (grp["Next_Catalyst_Date"].iloc[0] if grp["Next_Catalyst_Date"].notna().any() else "")
                    for col in ["Mechanism_of_Action", "Therapeutic_Area", "Market_Status"]:
                        if col in grp.columns:
                            v = grp[col].dropna()
                            row[col] = v.iloc[0] if len(v) else grp[col].iloc[0]
                    return pd.Series(row)
                assets = assets.groupby("_norm", as_index=False).apply(_merge_assets).reset_index(drop=True)
                assets = assets.drop(columns=["_norm"], errors="ignore")
                assets_display = assets.copy()
                drug_cols = ["Asset_Name", "Highest_Phase", "Market_Status", "Active_Trial_Count", "Next_Catalyst_Date", "Mechanism_of_Action", "Therapeutic_Area", "Detailed_Conditions", "Trial_NCTIds"]
                drug_cols = [c for c in drug_cols if c in assets_display.columns]
                assets_display = assets_display[drug_cols]
                st.markdown("<p class='section-title'>药物管线 · Drugs</p>", unsafe_allow_html=True)
                st.dataframe(assets_display.reset_index(drop=True), width="stretch", hide_index=True, height=min(400, 120 + len(assets_display) * 36), column_config={
                    "Asset_Name": st.column_config.TextColumn("Asset / 药物", width="medium"),
                    "Active_Trial_Count": st.column_config.NumberColumn("Trials", format="%d"),
                    "Next_Catalyst_Date": st.column_config.TextColumn("Catalyst", width="small"),
                    "Mechanism_of_Action": st.column_config.TextColumn("MoA", width="medium"),
                    "Therapeutic_Area": st.column_config.TextColumn("TA", width="medium"),
                    "Detailed_Conditions": st.column_config.TextColumn("Conditions", width="large"),
                    "Trial_NCTIds": st.column_config.TextColumn("NCT IDs", width="large"),
                })
            st.markdown("<p class='section-title'>试验 · Trials</p>", unsafe_allow_html=True)
            trials = df_trials_all[df_trials_all["Symbol"].astype(str).str.strip().str.upper() == selected_symbol.upper()].copy() if not df_trials_all.empty else pd.DataFrame()
            if trials.empty:
                st.info("该公司在 Enriched_Clinical_Trials.csv 中暂无试验记录。")
            else:
                trial_cols = ["NCTId", "Phases", "Status", "Conditions", "Interventions", "EnrollmentCount", "StartDate", "PrimaryCompletionDate", "BriefSummary", "OfficialTitle"]
                trial_cols = [c for c in trial_cols if c in trials.columns]
                trials_display = trials[trial_cols].copy()
                if "Interventions" in trials_display.columns:
                    trials_display["Interventions"] = trials_display["Interventions"].apply(_interventions_display)
                trials_display["NCT_Link"] = trials_display["NCTId"].apply(lambda x: f"https://clinicaltrials.gov/study/{x}" if pd.notna(x) and str(x).strip().upper().startswith("NCT") else "")
                st.dataframe(trials_display, width="stretch", hide_index=True, height=min(450, 120 + len(trials_display) * 38), column_config={
                    "NCTId": st.column_config.TextColumn("NCT ID", width="small"),
                    "NCT_Link": st.column_config.LinkColumn("CTG", display_text="Open", width="small"),
                    "EnrollmentCount": st.column_config.NumberColumn("Enrollment", format="%d"),
                    "BriefSummary": st.column_config.TextColumn("Summary", width="large"),
                    "OfficialTitle": st.column_config.TextColumn("Title", width="large"),
                })

        # ── Tab 2: 智能研报生成器（仅 Prompt）────────────────────────────────────
        with tab2:
            _company_row = df_summary[df_summary["Symbol"] == selected_symbol].iloc[0] if selected_symbol and not df_summary[df_summary["Symbol"] == selected_symbol].empty else None
            _mcap = _company_row.get("Market Cap") if _company_row is not None else None
            _cash = _company_row.get("Total Cash") if _company_row is not None else None
            _pipeline_rows = pd.DataFrame()
            if not df_master.empty and "Symbol" in df_master.columns:
                _pipeline_rows = df_master.loc[df_master["Symbol"].astype(str).str.strip().str.upper() == selected_symbol.upper()].copy()
            if not _pipeline_rows.empty and "Asset_Name" in _pipeline_rows.columns:
                _pipeline_rows = _pipeline_rows[~_pipeline_rows["Asset_Name"].apply(lambda x: _is_placebo_asset(x) if pd.notna(x) else False)]
            _dd_prompt = _build_dd_prompt(selected_symbol, company_name or "", _mcap or 0, _cash or 0, _pipeline_rows)
            st.markdown("**🧠 终极 DD Prompt（复制到 Gemini / Claude / Deep Research）**")
            _copy_id = f"prompt_hidden_{selected_symbol}"
            st.components.v1.html(
                f"""
                <textarea id="{_copy_id}" style="position:absolute;left:-9999px;top:-9999px;">{_dd_prompt.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")}</textarea>
                <button onclick="
                    var el=document.getElementById('{_copy_id}');
                    var txt=el.value;
                    navigator.clipboard.writeText(txt).then(function(){{
                        this.textContent='✅ 已复制';
                        var btn=this;
                        setTimeout(function(){{btn.textContent='📋 一键复制'}},2000);
                    }}.bind(this));
                " style="background:#2d5a87;color:white;border:none;border-radius:6px;padding:6px 16px;cursor:pointer;font-size:14px;">📋 一键复制</button>
                """,
                height=50,
            )
            st.text_area("Prompt 内容", value=_dd_prompt, height=500, key=f"dd_prompt_ta_{selected_symbol}")
            st.caption("复制上方 Prompt → 粘贴到 Gemini/Claude → 生成研报 → 回到 Tab 3 上传保存")

        # ── Tab 3: 研报上传、阅读与投研笔记 ──────────────────────────────────────
        with tab3:
            # ── 上传研报 ──────────────────────────────────────────────────────────
            st.markdown("**📤 上传研报**")
            _report_upload = st.file_uploader(
            "上传研报 (支持 .md / .txt / .docx / .doc / .rtf / .gdoc)",
            type=["md", "txt", "docx", "doc", "rtf", "gdoc"],
                key=f"report_uploader_{selected_symbol}",
            )
            if _report_upload is not None and selected_symbol:
                if st.button("💾 保存研报", type="primary", key=f"save_report_btn_{selected_symbol}"):
                    _dd_dir = _ensure_ai_dd_report_dir(selected_symbol)
                    _ts = datetime.now().strftime("%Y%m%d_%H%M%S")
                    _name = (_report_upload.name or "").lower()
                    if _name.endswith(".docx"):
                        _ext = ".docx"
                    elif _name.endswith(".doc"):
                        _ext = ".doc"
                    elif _name.endswith(".rtf"):
                        _ext = ".rtf"
                    elif _name.endswith(".gdoc"):
                        _ext = ".gdoc"
                    elif _name.endswith(".txt"):
                        _ext = ".txt"
                    else:
                        _ext = ".md"
                    _save_name = f"Report_{_ts}{_ext}"
                    _save_path = _dd_dir / _save_name
                    if _ext in (".doc", ".docx", ".rtf"):
                        _save_path.write_bytes(_report_upload.getvalue())
                    else:
                        _save_path.write_text(_report_upload.getvalue().decode("utf-8", errors="replace"), encoding="utf-8")
                    st.success(f"研报已保存至 {_save_path.relative_to(DATA_DIR)}")
                    st.rerun()
            st.markdown("---")
            # ── 历史研报阅读 ───────────────────────────────────────────────────────
            _report_files = _list_report_files(selected_symbol)
            if not _report_files:
                st.info("尚无该公司的历史研报，请先在上方上传。")
            else:
                _opts = [f"{human} - {fname}" for fname, human in _report_files]
                _sel = st.selectbox("选择历史研报", options=_opts, key=f"report_selector_{selected_symbol}")
                if _sel:
                    _sel_fname = _sel.split(" - ", 1)[1] if " - " in _sel else _report_files[0][0]
                    _report_path = AI_DD_REPORT_DIR / selected_symbol.upper() / _sel_fname
                    _report_content = ""
                    _parse_error = None
                    _suffix = _report_path.suffix.lower()
                    if _report_path.exists():
                        _report_content, _parse_error = _extract_any_text(_report_path)
                    _is_gdoc = _suffix == ".gdoc"
                    # .doc 解析成功时直接显示文本；失败时提供下载
                    _is_doc = _suffix == ".doc" and not _report_content.strip()
                    _is_docx = _suffix == ".docx" and not _report_content.strip()
                    # 笔记：每家公司共用一份永久笔记 Notes_{TICKER}.md，无论切换哪篇研报都加载同一文件
                    _company_notes_fname = f"Notes_{selected_symbol.upper()}.md"
                    _company_notes_path = AI_DD_REPORT_DIR / selected_symbol.upper() / _company_notes_fname
                    _company_notes_path.parent.mkdir(parents=True, exist_ok=True)
                    # 从磁盘读取（每次渲染重新读，保证不同标签页间同步）
                    _notes_content = ""
                    if _company_notes_path.exists():
                        _notes_content = _company_notes_path.read_text(encoding="utf-8", errors="replace")
                    # AI Summary：按了再 Summary，不按不调用
                    _summary_key = f"ai_summary_{selected_symbol}_{_sel_fname}"
                    if st.button("🤖 再 Summary", key=f"ai_summary_btn_{selected_symbol}_{_sel_fname}", help="调用 Gemini Flash 对当前研报生成简明总结"):
                        with st.spinner("正在调用 Gemini 生成总结…"):
                            _sum, _sum_err = _gemini_summarize_report(_report_content)
                        if _sum_err:
                            st.error(f"总结生成失败：{_sum_err}")
                        elif _sum:
                            st.session_state[_summary_key] = _sum
                            st.success("AI 总结已生成。")
                    if st.session_state.get(_summary_key):
                        with st.expander("📋 AI 总结", expanded=True):
                            st.markdown(st.session_state[_summary_key])
                    col1, col2 = st.columns([2, 1])
                    with col1:
                        st.markdown("**📄 研报内容**")
                        if _is_gdoc:
                            # .gdoc：提取 URL 显示链接
                            _gdoc_url = None
                            try:
                                _gdoc_data = json.loads(_report_path.read_text(encoding="utf-8", errors="replace"))
                                if isinstance(_gdoc_data, dict):
                                    _gdoc_url = _gdoc_data.get("url")
                            except Exception:
                                pass
                            if _gdoc_url:
                                st.link_button("📄 在 Google Docs 中打开", _gdoc_url)
                                st.caption("Google Docs 链接，点击上方按钮查看全文。")
                            else:
                                st.warning("无法从 .gdoc 文件中提取链接。")
                        elif _report_content.strip():
                            # 所有能成功提取文本的格式（.docx/.doc/.rtf/.md/.txt）统一显示
                            try:
                                _report_container = st.container(height=600)
                                with _report_container:
                                    st.markdown(_report_content)
                            except TypeError:
                                st.markdown(_report_content)
                        else:
                            # 解析失败：提供下载
                            if _parse_error:
                                st.warning(f"解析失败：{_parse_error}")
                            _mime_map = {
                                ".doc": "application/msword",
                                ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                ".rtf": "application/rtf",
                            }
                            _dl_mime = _mime_map.get(_suffix, "application/octet-stream")
                            st.download_button(
                                f"📥 下载 {_suffix.upper()} 文件",
                                data=_report_path.read_bytes(),
                                file_name=_report_path.name,
                                mime=_dl_mime,
                                key=f"dl_fallback_{selected_symbol}_{_sel_fname}",
                            )
                            st.caption("无法在此解析正文，请下载后用对应软件查看。")
                    with col2:
                        st.markdown(f"**✍️ 我的投研笔记 — {selected_symbol.upper()}**")
                        st.caption(f"文件：{_company_notes_fname}（切换研报不影响笔记，持续累积）")
                        _notes_ta = st.text_area(
                            "随手记录（可持续追加）",
                            value=_notes_content,
                            height=500,
                            key=f"notes_ta_{selected_symbol}",
                        )
                        if st.button("💾 保存笔记", type="primary", key=f"save_notes_{selected_symbol}"):
                            _company_notes_path.write_text(_notes_ta, encoding="utf-8")
                            st.success("笔记已保存。")
                            st.rerun()

# ── Portfolio Optimizer Tab ────────────────────────────────────────
with _tab_portfolio:
    st.caption("基于你的全局研报知识库，生成资产配置 Prompt，交由 Gemini Deep Research 生成组合方案。")
    st.markdown("---")

    # ── 1. 扫描知识库，构建 Master Archive ───────────────────────────
    st.markdown("### 📚 Step 1：全局投研知识库扫描")

    def _build_master_archive() -> tuple[str, list[str]]:
        """
        遍历 AI_DD_REPORT/{Symbol}/ 下的 Notes_{Symbol}.md 和
        Report_*.md/.txt/.docx，拼接成结构化长文本。
        返回 (archive_text, symbols_found)。
        """
        if not AI_DD_REPORT_DIR.exists():
            return "", []

        # 所有支持提取文本的格式
        _ARCHIVE_SUFFIXES = {".md", ".txt", ".docx", ".doc", ".rtf", ".gdoc"}

        def _read_report_file(rf: pathlib.Path) -> str:
            """用通用提取器读取研报，所有支持格式均可。"""
            content, err = _extract_any_text(rf)
            if content and content.strip():
                return content.strip()
            return f"（解析失败：{err}）" if err else "（内容为空）"

        sections = []
        symbols_found = []
        for sym_dir in sorted(AI_DD_REPORT_DIR.iterdir()):
            if not sym_dir.is_dir():
                continue
            sym = sym_dir.name.upper()
            notes_text = ""
            report_texts = []

            # 笔记：Notes_{TICKER}.md
            notes_path = sym_dir / f"Notes_{sym}.md"
            if notes_path.exists():
                notes_text = notes_path.read_text(encoding="utf-8", errors="replace").strip()

            # 研报：支持所有文本格式，按文件名倒序（最新在前）
            report_files = sorted(
                [f for f in sym_dir.iterdir()
                 if f.is_file()
                 and f.name.startswith("Report_")
                 and f.suffix.lower() in _ARCHIVE_SUFFIXES],
                key=lambda f: f.name, reverse=True,
            )
            for rf in report_files:
                content = _read_report_file(rf)
                if content:
                    ts_label = rf.stem[7:] if rf.stem.startswith("Report_") else rf.stem
                    try:
                        dt = datetime.strptime(ts_label, "%Y%m%d_%H%M%S")
                        ts_label = dt.strftime("%Y-%m-%d %H:%M")
                    except ValueError:
                        pass
                    fmt = rf.suffix.upper().lstrip(".")
                    report_texts.append(f"### 研报（{ts_label} · {fmt}）\n\n{content}")

            if not notes_text and not report_texts:
                continue
            symbols_found.append(sym)
            block = f"""
# ==========================================
# [公司代码: {sym}] 深度尽调与核心笔记
# ==========================================

## 我的投研笔记 (Notes)

{notes_text if notes_text else "（暂无笔记）"}

## AI 深度研报 (Report)

{chr(10).join(report_texts) if report_texts else "（暂无研报）"}
"""
            sections.append(block.strip())
        archive = "\n\n".join(sections)
        return archive, symbols_found

    _archive_text, _archive_symbols = _build_master_archive()

    if not _archive_text:
        st.warning("⚠️ `AI_DD_REPORT/` 目录为空或暂无文本研报。请先在各公司的 Tab 3 中上传研报或写入笔记。")
    else:
        _col_info, _col_dl = st.columns([3, 1])
        with _col_info:
            st.success(
                f"✅ 已扫描到 **{len(_archive_symbols)}** 家公司的研报/笔记：`{'`、`'.join(_archive_symbols)}`"
            )
            st.caption(f"总字符数：{len(_archive_text):,} 字符（约 {len(_archive_text)//4:,} tokens）")
        with _col_dl:
            st.download_button(
                label="📥 下载全局知识库",
                data=_archive_text.encode("utf-8"),
                file_name="Master_DD_Archive.md",
                mime="text/markdown",
                help="下载后作为附件上传到 Gemini Deep Research，配合下方 Prompt 使用",
            )

    st.markdown("---")

    # ── 2. 资产配置参数输入 ───────────────────────────────────────────
    st.markdown("### ⚙️ Step 2：配置投资参数")

    _param_c1, _param_c2 = st.columns(2)
    with _param_c1:
        _total_capital = st.number_input(
            "💰 计划投资总金额 ($)",
            min_value=1000,
            max_value=10_000_000,
            value=10_000,
            step=1000,
            format="%d",
            key="po_capital",
        )
        _horizon = st.selectbox(
            "🎯 投资周期预期",
            options=[
                "未来 6 个月（博弈核心催化剂）",
                "1-2 年（等待管线读出与商业化）",
                "3 年以上（长期价值回归）",
            ],
            index=0,
            key="po_horizon",
        )
    with _param_c2:
        _all_po_symbols = _archive_symbols if _archive_symbols else sorted(
            df_summary["Symbol"].dropna().astype(str).str.strip().str.upper().unique().tolist()
        )
        _selected_symbols = st.multiselect(
            "🏢 参与优化的公司（可剔除不想纳入的）",
            options=_all_po_symbols,
            default=_all_po_symbols,
            key="po_symbols",
        )
        _risk_notes = st.text_area(
            "📝 额外风险偏好或约束（可选）",
            placeholder="例如：不持有任何 Phase 1 管线；单票不超过 20%；回避 CNS 赛道……",
            height=100,
            key="po_extra_notes",
        )

    st.markdown("---")

    _gen_prompt_btn = st.button("⚡ 生成资产配置 Prompt", type="primary", key="po_gen_prompt")
    if _gen_prompt_btn:
        st.session_state["po_prompt_generated"] = True

    st.markdown("---")

    # ── 3. 生成 Gemini Deep Research Prompt ──────────────────────────
    st.markdown("### 🤖 Step 3：生成资产配置 Prompt")

    if not st.session_state.get("po_prompt_generated"):
        st.info("请在上方配置好参数后点击 **⚡ 生成资产配置 Prompt** 按钮。")
    elif not _selected_symbols:
        st.warning("请先在上方选择至少一家公司。")
    else:
        _syms_str = "、".join(_selected_symbols)
        _extra_block = f"\n额外约束：{_risk_notes.strip()}" if _risk_notes and _risk_notes.strip() else ""

        _po_prompt = f"""你是一位资深 Biotech 基金经理，同时精通量化风险模型与临床试验统计学。

我已经为你提供了【全局投研知识库】作为附件（Master_DD_Archive.md）。该文件包含了我针对以下公司的深度尽调报告（DD Report）与个人核心定性笔记（Notes）：
{_syms_str}

请仔细阅读该附件的全部内容，结合你通过 Deep Research 检索到的最新市场动态、临床试验进展、SEC 文件和分析师报告，为我生成一份投资组合优化方案。
请在给出最终投资组合之前，先用一段话简要总结你从知识库中提取到的关于这些标的的核心多空因素。

【投资参数】
- 总资金：${_total_capital:,}
- 投资周期：{_horizon}
- 标的池：仅限 {_syms_str}（允许配置现金 CASH）{_extra_block}

【核心资产配置约束原则】
1. 机制去相关性（MoA Diversification）：绝不将高比例资金集中于单一作用机制或高度相关靶点，防范一损俱损的系统性风险。
2. 资金生死线严控：惩罚 Cash Runway 不足 12 个月且面临重大定增砸盘风险的标的权重。
3. 赔率与胜率加权：优先参考我的"个人笔记"，基于我主观判定的临床试验水分、统计学胜率及当前市值赔率进行仓位倾斜。
4. 差异化仓位管理（Position Sizing）：每个标的的仓位大小应根据各组合策略目标独立优化，不要求等权——高确定性标的可给予更高权重，高风险催化剂标的相应控制敞口，明确说明每笔仓位大小背后的决策逻辑。
5. 治疗领域平衡（TA Diversification，如适用）：在标的池条件允许的情况下，尽量避免将资金过度集中于单一治疗领域（如全部为 CNS 或全部为代谢病），以降低赛道级系统性风险；若标的池本身较为集中则说明原因并给出风险提示。

请为我构建以下 4 种投资组合，每种组合须给出：
① 各标的具体分配金额与资金比例（精确到美元）
② 是否留有现金及比例
③ 严密的统计学与逻辑支撑（引用知识库中的具体数据）
④ 主要风险与对冲建议

---

### 🛡️ 组合 1：防御型（Risk-Averse）
侧重已商业化、拥有宽广护城河和充裕现金流的标的。追求极低回撤，对抗生物制药板块的高 Beta 风险。适合市场动荡期的底仓配置。

### ⚖️ 组合 2：均衡型（Core-Satellite）
50% 资金配置高确定性底仓（已商业化或 Phase 3 后期），50% 分散配置于未来 12-18 个月内有重大 Phase 2/3 数据读出的高赔率管线。

### 🚀 组合 3：进攻型（Catalyst-Driven / High-Beta）
重点下注知识库中赔率极高、且机制最具 First-in-class 潜力的中小市值 Biotech，博弈即将到来的催化剂事件。单票集中度可提高，但须分散 MoA。

### 🎯 组合 4：并购博弈型（M&A Arbitrage）
结合全球 Top 20 Big Pharma 的近一年 BD/M&A 交易记录（靶点偏好、适应症、交易规模、溢价倍数），以及各 Pharma 当前管线的薄弱环节，判断标的池中哪些公司最可能成为收购目标。
请给出：
- 潜在买家预测（列出 1-2 家最可能的 Pharma）
- 收购时间窗口预测（12 个月内 / 1-2 年内 / 2 年以上）
- 预估收购溢价区间（相对当前市值的 %）
- 博弈逻辑支撑

---

请用中文输出完整报告，数据引用须注明来源（知识库 or Deep Research）。"""

        st.code(_po_prompt, language="markdown")

        # 一键复制按钮
        _po_copy_id = "po_prompt_hidden"
        st.components.v1.html(
            f"""
            <textarea id="{_po_copy_id}" style="position:absolute;left:-9999px;top:-9999px;">{_po_prompt.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")}</textarea>
            <button onclick="
                var el=document.getElementById('{_po_copy_id}');
                navigator.clipboard.writeText(el.value).then(function(){{
                    this.textContent='✅ 已复制';
                    var btn=this;
                    setTimeout(function(){{btn.textContent='📋 一键复制 Prompt'}},2500);
                }}.bind(this));
            " style="background:#2d5a87;color:white;border:none;border-radius:6px;padding:8px 20px;cursor:pointer;font-size:14px;margin-top:8px;">📋 一键复制 Prompt</button>
            <span style="margin-left:12px;color:#6b7280;font-size:13px;">复制后在 Gemini Deep Research 中粘贴，并上传 Master_DD_Archive.md 作为附件</span>
            """,
            height=55,
        )

# ══════════════════════════════════════════════════════════════════
# 📈 Strategy Arena Tab
# ══════════════════════════════════════════════════════════════════
with _tab_arena:

    # ── helpers ───────────────────────────────────────────────────

    def _pt_safe_name(name: str) -> str:
        """策略名 → 文件名安全字符串（替换非法字符）。"""
        safe = re.sub(r'[\\/:*?"<>|]', "_", name).strip()
        return safe or "unnamed"

    def _pt_save(portfolio: dict) -> pathlib.Path:
        """将策略持久化为 JSON。
        文件名格式：{safe_name}_{YYYYMMDD}.json
        同一策略同一天保存则覆盖；不同天则留下历史版本。
        返回写入路径。
        """
        name = portfolio.get("portfolio_name", "unnamed")
        safe = _pt_safe_name(name)
        today = datetime.now().strftime("%Y%m%d")
        path = PT_DIR / f"{safe}_{today}.json"
        portfolio.setdefault("saved_at", datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        path.write_text(json.dumps(portfolio, ensure_ascii=False, indent=2), encoding="utf-8")
        return path

    def _pt_load(stem: str) -> dict | None:
        """按文件 stem（含日期后缀）加载策略 JSON。"""
        path = PT_DIR / f"{stem}.json"
        if not path.exists():
            return None
        try:
            return json.loads(path.read_text(encoding="utf-8"))
        except Exception:
            return None

    def _pt_list_files() -> list[dict]:
        """扫描 PT_DIR，返回所有策略文件信息列表（按日期降序）。
        每项：{"stem": str, "name": str, "date": str, "label": str}
        stem 格式：{safe_name}_{YYYYMMDD}
        """
        result = []
        for f in PT_DIR.glob("*.json"):
            if f.name == "portfolio_history.csv":
                continue
            stem = f.stem
            # 尝试拆出日期后缀（最后 8 位全是数字）
            parts = stem.rsplit("_", 1)
            if len(parts) == 2 and parts[1].isdigit() and len(parts[1]) == 8:
                raw_name = parts[0]
                date_str = parts[1]
                date_label = f"{date_str[:4]}-{date_str[4:6]}-{date_str[6:]}"
            else:
                raw_name = stem
                # 无日期后缀：用文件修改时间作为 date_str（用于排序）
                try:
                    _mtime = datetime.fromtimestamp(f.stat().st_mtime)
                    date_str = _mtime.strftime("%Y%m%d")
                    date_label = _mtime.strftime("%Y-%m-%d") + "（文件时间）"
                except Exception:
                    date_str = "00000000"
                    date_label = "（日期未知）"
            # 尝试从 JSON 内部读 portfolio_name（原始未 sanitize 的名字）
            try:
                data = json.loads(f.read_text(encoding="utf-8"))
                display_name = data.get("portfolio_name", raw_name)
                saved_at = data.get("saved_at", "")
                if saved_at:
                    date_label = saved_at[:10]
            except Exception:
                display_name = raw_name
                saved_at = ""
            result.append({
                "stem": stem,
                "name": display_name,
                "date": date_str,
                "saved_at": saved_at,
                "label": f"{display_name}  [{date_label}]",
            })
        return sorted(result, key=lambda x: x["date"], reverse=True)

    def _pt_list_names() -> list[str]:
        """兼容旧接口：返回最新版本策略名列表（去重，每个策略只保留最新一天）。"""
        seen: dict[str, dict] = {}
        for item in _pt_list_files():
            n = item["name"]
            if n not in seen or item["date"] > seen[n]["date"]:
                seen[n] = item
        return [v["stem"] for v in sorted(seen.values(), key=lambda x: x["name"])]

    def _pt_list_unique_names() -> list[str]:
        """返回去重后的策略 portfolio_name 列表（不含日期后缀）。"""
        seen: set = set()
        result = []
        for item in _pt_list_files():
            if item["name"] not in seen:
                seen.add(item["name"])
                result.append(item["name"])
        return result

    def _fetch_prices(symbols: list[str]) -> dict[str, float]:
        """批量用 yfinance 获取最新收盘价，返回 {symbol: price}。"""
        prices = {}
        for sym in symbols:
            try:
                info = yf.Ticker(sym).info or {}
                p = info.get("currentPrice") or info.get("regularMarketPrice") or info.get("previousClose")
                prices[sym] = float(p) if p else 0.0
            except Exception:
                prices[sym] = 0.0
        return prices

    def _calc_portfolio_value(portfolio: dict, prices: dict[str, float]) -> float:
        total = portfolio.get("cash", 0.0)
        for h in portfolio.get("holdings", []):
            sym = h.get("symbol", "").upper()
            shares = float(h.get("shares", 0))
            total += shares * prices.get(sym, h.get("avg_cost", 0))
        return total

    def _pt_append_history(name: str, value: float) -> None:
        today = datetime.now().strftime("%Y-%m-%d")
        row = pd.DataFrame([{"Date": today, "Portfolio_Name": name, "Total_Value": round(value, 2)}])
        if PT_HISTORY_CSV.exists():
            hist = pd.read_csv(PT_HISTORY_CSV)
            # 同一天同一策略只保留最新值
            hist = hist[~((hist["Date"] == today) & (hist["Portfolio_Name"] == name))]
            hist = pd.concat([hist, row], ignore_index=True)
        else:
            hist = row
        hist.to_csv(PT_HISTORY_CSV, index=False, encoding="utf-8")

    def _pt_load_history(names: list[str]) -> pd.DataFrame:
        if not PT_HISTORY_CSV.exists():
            return pd.DataFrame(columns=["Date", "Portfolio_Name", "Total_Value"])
        hist = pd.read_csv(PT_HISTORY_CSV)
        return hist[hist["Portfolio_Name"].isin(names)].copy()

    # ─────────────────────────────────────────────────────────────
    st.markdown("## 📈 策略竞技场 (Strategy Arena & Paper Trading)")
    st.caption("创建、追踪、对比多个模拟策略组合，支持 AI 批量导入与历史回溯。")
    st.markdown("---")

    arena_tab_a, arena_tab_b, arena_tab_c = st.tabs([
        "⚡ A · 批量建仓 & AI 导入",
        "🏇 B · 多策略赛马对比",
        "🔍 C · 历史回溯归因",
    ])

    # ══════════════════════════════════════════════════════════════
    # MODULE A — 批量建仓 & AI 导入
    # ══════════════════════════════════════════════════════════════
    with arena_tab_a:
        st.markdown("### ⚡ AI 批量解析研报 → 生成策略草案")
        st.caption("总资金、持仓股数、现金由 AI 从研报中提取，解析后可在表格中自由修改，系统自动校验三者是否自洽。")

        # ── 研报来源：上传新文件 或 选取历史存档 ───────────────────────
        def _arena_list_saved_reports() -> list[pathlib.Path]:
            """扫描 Arena_Reports/ 并按时间戳降序返回文件列表。"""
            files = sorted(
                [f for f in PT_ARENA_REPORTS_DIR.iterdir()
                 if f.is_file() and f.name.startswith("Report_")],
                key=lambda f: f.name,
                reverse=True,
            )
            return files

        def _arena_report_label(p: pathlib.Path) -> str:
            """生成人类可读标签：文件名时间戳 → 日期，或回退到 mtime。"""
            # 格式：Report_YYYYMMDD_HHMMSS_{原文件名}.ext
            stem = p.stem
            if stem.startswith("Report_") and len(stem) >= 22:
                ts_part = stem[7:22]  # YYYYMMDD_HHMMSS
                try:
                    dt = datetime.strptime(ts_part, "%Y%m%d_%H%M%S")
                    orig = stem[23:] if len(stem) > 23 else ""
                    orig_label = f"  {orig}{p.suffix}" if orig else ""
                    return f"{dt.strftime('%Y-%m-%d %H:%M:%S')}{orig_label}"
                except ValueError:
                    pass
            mtime = datetime.fromtimestamp(p.stat().st_mtime).strftime("%Y-%m-%d %H:%M:%S")
            return f"{p.name}（修改于 {mtime}）"

        _saved_reports = _arena_list_saved_reports()

        _src_tab_new, _src_tab_saved = st.tabs(["📤 上传新研报", "📂 选取已保存研报"])

        _a_report_bytes: bytes | None = None
        _a_report_suffix: str = ".md"
        _a_report_source_name: str = ""

        with _src_tab_new:
            _a_col1, _a_col2 = st.columns([4, 1])
            with _a_col1:
                _a_report_file = st.file_uploader(
                    "上传研报文件（.md / .txt / .docx）",
                    type=["md", "txt", "docx"],
                    key="arena_report_upload",
                )
            with _a_col2:
                _a_api_key = st.text_input(
                    "Gemini API Key（留空则用 .env）",
                    value="", type="password", key="arena_api_key",
                )
            if _a_report_file is not None:
                _a_report_bytes = _a_report_file.getvalue()
                _a_report_suffix = pathlib.Path(_a_report_file.name).suffix.lower()
                _a_report_source_name = _a_report_file.name
                # 自动保存到 Arena_Reports/
                _save_ts = datetime.now().strftime("%Y%m%d_%H%M%S")
                _safe_orig = re.sub(r"[^\w\-_.]", "_", _a_report_file.name)[:60]
                _save_fname = f"Report_{_save_ts}_{_safe_orig}"
                _save_path = PT_ARENA_REPORTS_DIR / _save_fname
                if not _save_path.exists():
                    _save_path.write_bytes(_a_report_bytes)
                    st.caption(f"✅ 已自动存档：`{_save_fname}`")

        with _src_tab_saved:
            _a_api_key_saved = st.text_input(
                "Gemini API Key（留空则用 .env）",
                value="", type="password", key="arena_api_key_saved",
            )
            if not _saved_reports:
                st.info("暂无存档研报，请先在「上传新研报」标签页上传一份。")
            else:
                _saved_labels = [_arena_report_label(f) for f in _saved_reports]
                _saved_label_map = dict(zip(_saved_labels, _saved_reports))
                _sel_saved_label = st.selectbox(
                    f"选择已保存研报（共 {len(_saved_reports)} 份）",
                    options=_saved_labels,
                    key="arena_saved_report_sel",
                )
                _sel_saved_path = _saved_label_map[_sel_saved_label]
                _a_report_bytes = _sel_saved_path.read_bytes()
                _a_report_suffix = _sel_saved_path.suffix.lower()
                _a_report_source_name = _sel_saved_path.name
                # 可选：删除存档
                if st.button("🗑️ 删除此存档", key="arena_del_report_btn"):
                    _sel_saved_path.unlink(missing_ok=True)
                    st.success("已删除。")
                    st.rerun()
                st.caption(f"来源：`{_sel_saved_path.name}`")
                # 从历史选取时，api key 用本 tab 的输入
                if _a_api_key_saved.strip():
                    _a_api_key = _a_api_key_saved

        _AI_BATCH_PROMPT = """你现在是一个没有任何创造力的"无情的数据提取机器"。你的唯一任务是：精确提取用户上传的研报原文中【已经明确列出】的投资组合方案。

【🛑 绝对禁止事项 - 核心红线】：
1. 绝不允许你发挥任何创造力！绝不允许你自己编造、补充或推演任何新的投资组合！
2. 研报原文里明确写了几个组合，你就只能提取几个。如果原文只有 4 个组合，你的返回结果中就必须【有且仅有 4 个】组合。
3. 绝不允许添加原文中没有提到的股票代码 (Symbol)。

请严格按照以下 JSON 数组格式返回（绝对不要包含任何 Markdown 标记，如 ```json，直接输出中括号开头和结尾）：
[
  {
    "portfolio_name": "原文中的策略名称(如: 进攻型组合)",
    "total_capital": 200000,
    "cash": null,
    "holdings": [
      {"symbol": "NVO", "shares": null, "avg_cost": null, "allocation_pct": 40, "amount": 80000, "rationale": "提取原文中的一句话建仓理由"},
      {"symbol": "AMGN", "shares": null, "avg_cost": null, "allocation_pct": 60, "amount": 120000, "rationale": "..."}
    ]
  }
]

字段说明（从原文中能读到什么就填什么，读不到的填 null）：
- total_capital：原文中的总资金规模（数字）
- cash：原文中的现金仓位（数字）
- holdings[].symbol：股票代码（NASDAQ/NYSE 代码）
- holdings[].shares：持仓股数（若原文有）
- holdings[].avg_cost：建仓均价（若原文有）
- holdings[].allocation_pct：仓位占比百分比（若原文有，如 40 表示 40%）
- holdings[].amount：持仓金额（若原文有）
- holdings[].rationale：直接引用原文中的一句建仓理由

[用户上传的投研报告原文开始]：
{REPORT_TEXT}
[用户上传的投研报告原文结束]
"""

        if _a_report_bytes:
            st.caption(f"📄 当前研报：**{_a_report_source_name}**")

        if st.button("🤖 AI 解析研报 → 生成策略草案", key="arena_ai_parse_btn"):
            if _a_report_bytes is None:
                st.warning("请先上传研报文件，或从「选取已保存研报」中选择一份。")
            else:
                _raw = _a_report_bytes
                _suf = _a_report_suffix
                if _suf == ".docx":
                    _tmp = PT_DIR / "_tmp_upload.docx"
                    _tmp.write_bytes(_raw)
                    _report_txt, _ = _extract_docx_text(_tmp)
                    _tmp.unlink(missing_ok=True)
                else:
                    _report_txt = _raw.decode("utf-8", errors="replace")
                _use_key = _a_api_key.strip() or _GEMINI_API_KEY
                if not _use_key:
                    st.error("请输入 Gemini API Key 或在 .env 中配置 GEMINI_API_KEY。")
                else:
                    _cli = genai.Client(api_key=_use_key)
                    with st.spinner("正在调用 Gemini 生成策略草案（可能需要 30-60 秒）…"):
                        try:
                            _resp = _cli.models.generate_content(
                                model=GEMINI_FLASH_MODEL,
                                contents=_AI_BATCH_PROMPT.replace("{REPORT_TEXT}", _report_txt[:80000]),
                            )
                            _raw_json = re.sub(r"^```[a-z]*\n?", "", (_resp.text or "").strip())
                            _raw_json = re.sub(r"\n?```$", "", _raw_json).strip()
                            _parsed = json.loads(_raw_json)
                            if not isinstance(_parsed, list):
                                _parsed = [_parsed]
                            st.session_state["arena_parsed_portfolios"] = _parsed[:10]
                            st.success(f"✅ 成功解析出 {len(_parsed[:10])} 个策略草案，请在下方审核并修改。")
                        except json.JSONDecodeError as _je:
                            st.error(f"JSON 解析失败：{_je}")
                            st.code((_raw_json or "")[:2000])
                        except Exception as _e:
                            st.error(f"API 调用失败：{_e}")

        # ── 草案审核与手动编辑区 ──────────────────────────────────────
        _draft_portfolios = st.session_state.get("arena_parsed_portfolios", [])
        _all_valid = True  # 全部校验通过才允许部署

        if _draft_portfolios:
            st.markdown(f"---\n### 📋 草案审核（共 {len(_draft_portfolios)} 个策略）")
            _draft_hdr_col1, _draft_hdr_col2 = st.columns([4, 1])
            with _draft_hdr_col1:
                st.caption("可直接在表格中双击修改 Shares / Avg Cost。系统自动校验总资金是否自洽。")
            with _draft_hdr_col2:
                if st.button("🗑️ 清空全部草案", key="arena_clear_all_drafts"):
                    st.session_state["arena_parsed_portfolios"] = []
                    st.rerun()

            _edited_portfolios = []  # 存储编辑后准备部署的策略

            for _pi, _pf in enumerate(_draft_portfolios[:10]):
                _pf_name = _pf.get("portfolio_name", f"策略{_pi+1}")
                _title_col, _del_col = st.columns([8, 1])
                with _title_col:
                    st.markdown(f"#### {_pi+1}. {_pf_name}")
                with _del_col:
                    if st.button("✕ 删除", key=f"arena_del_draft_{_pi}", help=f"删除草案「{_pf_name}」"):
                        _new_drafts = [p for i, p in enumerate(_draft_portfolios) if i != _pi]
                        st.session_state["arena_parsed_portfolios"] = _new_drafts
                        st.rerun()

                # ── 名称可修改 ──
                _new_name = st.text_input(
                    "策略名称", value=_pf_name,
                    key=f"pf_name_{_pi}",
                )

                # ── 持仓表格可编辑 ──
                _holdings = _pf.get("holdings", [])

                # 找出需要拉价格的标的：avg_cost 为 null/0
                _syms_need_price = list({
                    (h.get("symbol") or "").upper()
                    for h in _holdings
                    if (h.get("symbol") or "").strip()
                    and not (h.get("avg_cost") or 0)
                })
                _live_prices: dict = {}
                if _syms_need_price:
                    with st.spinner(f"拉取 {', '.join(_syms_need_price)} 最新股价…"):
                        _live_prices = _fetch_prices(_syms_need_price)

                def _h_to_row(h: dict, live_px: dict) -> dict:
                    _sym = (h.get("symbol") or "").upper()
                    _shares_raw = h.get("shares")
                    _cost_raw = h.get("avg_cost")
                    _amount_raw = h.get("amount")
                    _alloc = h.get("allocation_pct")

                    _cost = float(_cost_raw) if _cost_raw else live_px.get(_sym, 0.0)
                    _amount = float(_amount_raw) if _amount_raw else None

                    if _shares_raw and float(_shares_raw) > 0:
                        # 有股数 → 直接用，amount 根据股数*价格重算
                        _shares = int(round(float(_shares_raw)))
                        _actual_amount = _shares * _cost
                    elif _amount and _cost > 0:
                        # 有金额无股数 → 计算整数股数，重算金额
                        _shares = int(_amount / _cost)
                        _actual_amount = _shares * _cost
                    else:
                        _shares = 0
                        _actual_amount = 0.0

                    return {
                        "Symbol": _sym,
                        "Shares": _shares,
                        "Avg Cost ($)": round(_cost, 4),
                        "Amount ($)": round(_actual_amount, 2),
                        "Alloc %": float(_alloc) if _alloc is not None else 0.0,
                        "Rationale": h.get("rationale", ""),
                    }

                _rows = [_h_to_row(h, _live_prices) for h in _holdings]
                _h_df = pd.DataFrame(_rows)

                # ── 资金校验：先确定总资金/CASH，以便正确计算 Alloc % ──
                _ai_total = _pf.get("total_capital")
                _stock_value_pre = _h_df["Amount ($)"].sum()

                # 总资金输入框（留 0 = 自动 = 纯股票市值）
                _override_total = st.number_input(
                    "总资金 $ (留 0 = 自动 = 纯股票市值)",
                    value=float(_ai_total) if _ai_total else 0.0,
                    min_value=0.0, step=1000.0, format="%.2f",
                    key=f"pf_total_{_pi}",
                )
                _use_total = _override_total if _override_total > 0 else None
                _pre_total = _use_total if _use_total else _stock_value_pre
                _pre_cash = max(0.0, (_use_total - _stock_value_pre) if _use_total else 0.0)

                # Alloc % 以含 CASH 的总资金为分母
                if _pre_total > 0:
                    _h_df["Alloc %"] = (_h_df["Amount ($)"] / _pre_total * 100).round(1)

                # ── 追加 CASH / TOTAL 行，合成一张完整表格 ──
                # 用哨兵标记列区分汇总行，避免 emoji 字符在 data_editor 中被改动后过滤失效
                _cash_alloc_pre = round(_pre_cash / _pre_total * 100, 1) if _pre_total > 0 else 0.0
                _h_df["_is_summary"] = False
                _footer_df = pd.DataFrame([
                    {"Symbol": "💵 CASH", "Shares": 0, "Avg Cost ($)": 0.0,
                     "Amount ($)": _pre_cash, "Alloc %": _cash_alloc_pre,
                     "Rationale": "", "_is_summary": True},
                    {"Symbol": "📊 TOTAL", "Shares": 0, "Avg Cost ($)": 0.0,
                     "Amount ($)": _pre_total, "Alloc %": 100.0,
                     "Rationale": "", "_is_summary": True},
                ])
                _full_df = pd.concat([_h_df, _footer_df], ignore_index=True)
                # 确保数值列类型一致，防止 concat 后 dtype=object 导致 data_editor 异常
                for _col in ["Shares", "Avg Cost ($)", "Amount ($)", "Alloc %"]:
                    _full_df[_col] = pd.to_numeric(_full_df[_col], errors="coerce").fillna(0.0)

                # 一张 data_editor；_is_summary 列隐藏，保存时靠它过滤
                _edited_full_df = st.data_editor(
                    _full_df.drop(columns=["_is_summary"]),
                    num_rows="dynamic",
                    use_container_width=True,
                    key=f"pf_holdings_editor_{_pi}",
                    column_config={
                        "Symbol": st.column_config.TextColumn("Symbol", width="small"),
                        "Shares": st.column_config.NumberColumn("股数 (整数)", format="%d"),
                        "Avg Cost ($)": st.column_config.NumberColumn("最新/均价 ($)", format="$%.4f"),
                        "Amount ($)": st.column_config.NumberColumn("持仓金额 ($)", format="$%.2f"),
                        "Alloc %": st.column_config.NumberColumn("Alloc % (占总资金)", format="%.1f%%"),
                        "Rationale": st.column_config.TextColumn("Rationale", width="large"),
                    },
                )

                # 过滤出真实持仓行（排除 CASH / TOTAL 汇总行，用原始 _is_summary 标记）
                _summary_mask = _full_df["_is_summary"].values
                # _edited_full_df 行数可能因用户新增行而变化，按 Symbol 过滤更稳健
                _SUMMARY_SYMS = {"💵 CASH", "📊 TOTAL"}
                _edited_h_df = _edited_full_df[
                    ~_edited_full_df["Symbol"].astype(str).str.strip().isin(_SUMMARY_SYMS)
                ]

                # ── 资金校验（用编辑后的 Amount 重算）──
                _stock_value = sum(
                    float(r.get("Amount ($)", 0))
                    for _, r in _edited_h_df.iterrows()
                    if str(r.get("Symbol", "")).strip() and float(r.get("Amount ($)", 0)) > 0
                )

                _valid = True
                _err_msg = ""
                if _use_total is None:
                    _final_total = _stock_value
                    _final_cash = 0.0
                elif _use_total >= _stock_value:
                    _final_total = _use_total
                    _final_cash = _use_total - _stock_value
                else:
                    _final_total = _use_total
                    _final_cash = 0.0
                    _valid = False
                    _err_msg = (
                        f"⚠️ 股票持仓市值 ${_stock_value:,.0f} 超过了总资金 ${_use_total:,.0f}，"
                        f"请减少持仓数量或提高总资金。"
                    )

                if not _valid:
                    st.error(_err_msg)
                    _all_valid = False
                else:
                    st.success(f"✅ 资金自洽：股票 ${_stock_value:,.0f} + 现金 ${_final_cash:,.0f} = 总资金 ${_final_total:,.0f}")

                # 分配图（含 CASH）
                _bar_data = {r["Symbol"]: float(r.get("Amount ($)", 0)) for _, r in _edited_h_df.iterrows() if float(r.get("Amount ($)", 0)) > 0}
                if _final_cash and _final_cash > 0:
                    _bar_data["💵 CASH"] = _final_cash
                if _bar_data:
                    st.bar_chart(pd.DataFrame({"市值 ($)": _bar_data}), height=180)

                st.markdown("---")

                # 收集编辑后数据（Alloc % 以总资金含 CASH 为分母）
                _edited_holdings = []
                for _, _hr in _edited_h_df.iterrows():
                    _sym = str(_hr.get("Symbol", "")).strip().upper()
                    if not _sym:
                        continue
                    _amt = float(_hr.get("Amount ($)", 0))
                    _alloc_pct = round(_amt / _final_total * 100, 1) if _final_total > 0 else 0.0
                    _edited_holdings.append({
                        "symbol": _sym,
                        "shares": int(_hr.get("Shares", 0)),
                        "avg_cost": float(_hr.get("Avg Cost ($)", 0)),
                        "amount": _amt,
                        "allocation_pct": _alloc_pct,
                        "rationale": str(_hr.get("Rationale", "")),
                    })
                _edited_portfolios.append({
                    "portfolio_name": _new_name,
                    "total_capital": _final_total or 0,
                    "cash": _final_cash or 0,
                    "holdings": _edited_holdings,
                    "_valid": _valid,
                })

            # ── 批量部署按钮 ──
            if not _all_valid:
                st.warning("⚠️ 有策略资金校验未通过，请修正后再部署。")
            if st.button(
                "⚡ 批量部署至虚拟盘仓库",
                type="primary", key="arena_deploy_btn",
                disabled=not _all_valid,
            ):
                _deployed = 0
                for _ep in _edited_portfolios:
                    if not _ep.get("_valid", True):
                        continue
                    _ep_out = {k: v for k, v in _ep.items() if k != "_valid"}
                    _ep_out["created_at"] = datetime.now().strftime("%Y-%m-%d")
                    _pt_save(_ep_out)
                    _deployed += 1
                st.success(f"✅ 已成功部署 {_deployed} 个策略至 `Paper_Trading_Portfolios/`！")
                st.rerun()

        st.markdown("---")
        st.markdown("### ✍️ 手动建仓（加入草案审核）")
        st.caption(
            "每行格式支持以下写法：  \n"
            "`ACAD 500 18.50` → 代码 股数 均价 &nbsp;|&nbsp; "
            "`ACAD 500` → 代码 股数（自动抓最新价）&nbsp;|&nbsp; "
            "`ACAD $9250` → 代码 金额（自动抓价，取整股数）"
        )
        with st.form("manual_portfolio_form"):
            _m_name = st.text_input("策略名称", value="我的自定义策略", key="manual_pf_name")
            _m_total = st.number_input(
                "总资金 $ (留 0 = 自动 = 纯股票市值；填入后 CASH 自动推算)",
                value=0.0, step=1000.0, format="%.2f", key="manual_pf_total",
            )
            _m_holdings_raw = st.text_area(
                "持仓（每行：代码 股数 [均价] 或 代码 $金额）",
                height=180, key="manual_holdings_raw",
                placeholder="ACAD 500 18.50\nALKS 300\nARWR $12000",
            )
            _m_submit = st.form_submit_button("➕ 加入草案审核区", type="primary")

        if _m_submit:
            _m_lines = [l.strip() for l in _m_holdings_raw.strip().splitlines() if l.strip()]

            # 判断哪些 symbol 需要拉价格
            _m_need_price_syms: set = set()
            _m_parsed_lines: list = []
            for _ml in _m_lines:
                _parts = _ml.split()
                if len(_parts) < 2:
                    continue
                _sym = _parts[0].upper()
                _val_raw = _parts[1]
                _is_amount = _val_raw.startswith("$")
                if _is_amount:
                    _m_need_price_syms.add(_sym)
                    _m_parsed_lines.append(("amount", _sym, float(_val_raw.lstrip("$")), None))
                elif len(_parts) >= 3:
                    _m_parsed_lines.append(("shares_price", _sym, float(_parts[1]), float(_parts[2])))
                else:
                    _m_need_price_syms.add(_sym)
                    _m_parsed_lines.append(("shares_only", _sym, float(_parts[1]), None))

            with st.spinner(f"拉取最新价格：{', '.join(_m_need_price_syms) or '无需'}…"):
                _m_prices = _fetch_prices(list(_m_need_price_syms)) if _m_need_price_syms else {}

            _m_holdings_out: list = []
            for _mode, _sym, _val, _extra in _m_parsed_lines:
                _px = _extra if _extra else _m_prices.get(_sym, 0.0)
                if _mode == "amount":
                    _sh = int(_val / _px) if _px > 0 else 0
                else:
                    _sh = int(round(_val))
                _actual_amt = round(_sh * _px, 2)
                _m_holdings_out.append({
                    "symbol": _sym, "shares": _sh, "avg_cost": round(_px, 4),
                    "amount": _actual_amt, "allocation_pct": 0, "rationale": "",
                })

            _m_stock_val = sum(h["amount"] for h in _m_holdings_out)
            _m_use_total = _m_total if _m_total > 0 else None

            # 自动推算 CASH（总资金 - 股票市值）
            if _m_use_total is None:
                _m_final_total = _m_stock_val
                _m_final_cash = 0.0
            elif _m_use_total >= _m_stock_val:
                _m_final_total = _m_use_total
                _m_final_cash = _m_use_total - _m_stock_val
            else:
                st.error(
                    f"⚠️ 股票市值 ${_m_stock_val:,.0f} 超过总资金 ${_m_use_total:,.0f}，"
                    "请减少持仓数量或提高总资金。未加入草案。"
                )
                st.stop()

            # 构建草案格式（与 AI 解析格式一致）并加入 session_state
            _m_draft = {
                "portfolio_name": _m_name,
                "total_capital": _m_final_total,
                "cash": _m_final_cash,
                "holdings": _m_holdings_out,
            }
            _existing_drafts = st.session_state.get("arena_parsed_portfolios", [])
            if len(_existing_drafts) >= 10:
                st.warning("⚠️ 草案区已满 10 个策略，请先部署或清除部分草案。")
            else:
                st.session_state["arena_parsed_portfolios"] = _existing_drafts + [_m_draft]
                st.success(
                    f"✅ 策略「{_m_name}」已加入草案审核区！"
                    f"总资金 ${_m_final_total:,.0f}（股票 ${_m_stock_val:,.0f} + 现金 ${_m_final_cash:,.0f}）"
                )
                st.rerun()

        # ── 删除已保存策略（仓库管理，放 Tab A 最底部）────────────────
        st.markdown("---")
        st.markdown("### 🗑️ 删除已保存策略（仓库管理）")
        _del_all_files = _pt_list_files()
        if not _del_all_files:
            st.caption("仓库为空，暂无可删除的策略。")
        else:
            _del_labels = [f["label"] for f in _del_all_files]
            _del_label_to_stem = {f["label"]: f["stem"] for f in _del_all_files}
            _del_selected = st.multiselect(
                "选择要删除的策略版本（可多选，删除后无法恢复）",
                options=_del_labels,
                key="arena_del_pf_select",
            )
            if _del_selected:
                if st.button(
                    f"🗑️ 确认删除所选 {len(_del_selected)} 个策略文件",
                    type="primary", key="arena_del_pf_btn",
                ):
                    _del_count = 0
                    for _dl in _del_selected:
                        _stem = _del_label_to_stem.get(_dl, "")
                        _path = PT_DIR / f"{_stem}.json"
                        if _path.exists():
                            _path.unlink()
                            _del_count += 1
                    st.success(f"✅ 已删除 {_del_count} 个策略文件。")
                    st.rerun()

    # ══════════════════════════════════════════════════════════════
    # MODULE B — 多策略赛马对比台
    # ══════════════════════════════════════════════════════════════
    with arena_tab_b:
        st.markdown("### 🏇 多策略实时赛马对比（最多 10 个）")
        _all_pf_files_b = _pt_list_files()
        if not _all_pf_files_b:
            st.info("仓库为空，请先在 Tab A 创建或导入策略。")
        else:
            # label → stem 映射（label 含策略名+保存日期）
            _b_label_to_stem = {f["label"]: f["stem"] for f in _all_pf_files_b}
            _b_labels = [f["label"] for f in _all_pf_files_b]
            _default_labels = _b_labels[:min(3, len(_b_labels))]
            _selected_labels = st.multiselect(
                "🔍 选择要追踪与对比的策略（最多 10 个）",
                options=_b_labels,
                default=_default_labels,
                max_selections=10,
                key="arena_selected_pfs",
            )
            _selected_pf_stems = [_b_label_to_stem[l] for l in _selected_labels]

            if not _selected_pf_stems:
                st.info("请选择至少一个策略。")
            else:
                if st.button("🔄 刷新实时净值", type="primary", key="arena_refresh_btn"):
                    _b_all_syms = []
                    _b_pf_data = {}
                    for _stem in _selected_pf_stems:
                        _pf = _pt_load(_stem)
                        if _pf:
                            _b_pf_data[_stem] = _pf
                            _b_all_syms += [h["symbol"].upper() for h in _pf.get("holdings", [])]
                    with st.spinner("拉取最新价格…"):
                        _b_prices = _fetch_prices(list(set(_b_all_syms)))
                    st.session_state["arena_b_prices"] = _b_prices
                    st.session_state["arena_b_pf_data"] = _b_pf_data
                    for _stem, _pf in _b_pf_data.items():
                        _val = _calc_portfolio_value(_pf, _b_prices)
                        _pt_append_history(_pf.get("portfolio_name", _stem), _val)
                    st.success("净值已更新并写入历史记录。")

                _b_prices = st.session_state.get("arena_b_prices", {})
                _b_pf_data = st.session_state.get("arena_b_pf_data", {})
                if not _b_pf_data:
                    for _stem in _selected_pf_stems:
                        _pf = _pt_load(_stem)
                        if _pf:
                            _b_pf_data[_stem] = _pf

                _metric_cols = st.columns(min(len(_selected_pf_stems), 5))
                for _ci, _stem in enumerate(_selected_pf_stems):
                    _pf = _b_pf_data.get(_stem) or _pt_load(_stem)
                    if not _pf:
                        continue
                    _cap = float(_pf.get("total_capital", 0))
                    _cur_val = _calc_portfolio_value(_pf, _b_prices) if _b_prices else _cap
                    _saved_at = _pf.get("saved_at", _pf.get("created_at", ""))[:10]
                    _today_str = datetime.now().strftime("%Y-%m-%d")
                    if _saved_at == _today_str or _cap == 0:
                        _pnl_str = "0.00%"
                        _delta_color = "off"
                    else:
                        _pnl = (_cur_val - _cap) / _cap * 100
                        _pnl_str = f"{_pnl:+.2f}%"
                        _delta_color = "normal"
                    _display_name = _pf.get("portfolio_name", _stem)
                    with _metric_cols[_ci % 5]:
                        st.metric(
                            label=_display_name,
                            value=f"${_cur_val:,.0f}",
                            delta=_pnl_str,
                            delta_color=_delta_color,
                        )
                        st.caption(f"初始资金 ${_cap:,.0f}　保存于 {_saved_at}")

                # 净值历史曲线（用 portfolio_name 作 key）
                _pf_names_for_hist = [
                    (_b_pf_data.get(s) or _pt_load(s) or {}).get("portfolio_name", s)
                    for s in _selected_pf_stems
                ]
                _hist = _pt_load_history(_pf_names_for_hist)
                if not _hist.empty:
                    st.markdown("**📊 净值对比曲线（含 XBI / SPY 基准）**")
                    _hist["Date"] = pd.to_datetime(_hist["Date"])
                    _hist_pivot = _hist.pivot_table(
                        index="Date", columns="Portfolio_Name",
                        values="Total_Value", aggfunc="last",
                    )
                    # ── 拉取 XBI / SPY 作 benchmark，按各策略最早日期起点归一化 ──
                    try:
                        _b_hist_start = _hist_pivot.index.min()
                        _b_hist_end = _hist_pivot.index.max()
                        if _b_hist_start < _b_hist_end:
                            _bm_raw = yf.download(
                                ["XBI", "SPY"],
                                start=_b_hist_start.strftime("%Y-%m-%d"),
                                end=(_b_hist_end + pd.Timedelta(days=1)).strftime("%Y-%m-%d"),
                                auto_adjust=True, progress=False,
                            )
                            if isinstance(_bm_raw.columns, pd.MultiIndex):
                                _bm_close = _bm_raw["Close"]
                            else:
                                _bm_close = _bm_raw[["Close"]]
                                _bm_close.columns = ["XBI"]
                            _bm_close.index = pd.to_datetime(_bm_close.index)
                            # 每个策略各自按初始资金归一化
                            for _bm_sym in [c for c in ["XBI", "SPY"] if c in _bm_close.columns]:
                                _bm_series = _bm_close[_bm_sym].dropna()
                                # 取所有策略的平均初始资金作为 benchmark 起始资金
                                _avg_cap = sum(
                                    float((_b_pf_data.get(s) or _pt_load(s) or {}).get("total_capital", 0))
                                    for s in _selected_pf_stems
                                ) / max(len(_selected_pf_stems), 1)
                                if len(_bm_series) > 0 and float(_bm_series.iloc[0]) > 0:
                                    _bm_norm = _bm_series / float(_bm_series.iloc[0]) * _avg_cap
                                    _bm_norm.name = f"{_bm_sym} (基准)"
                                    _hist_pivot = _hist_pivot.join(_bm_norm, how="left")
                    except Exception:
                        pass  # benchmark 拉取失败不影响主曲线
                    st.line_chart(_hist_pivot, height=380)
                else:
                    st.info("暂无历史净值数据，点击「刷新实时净值」后开始记录。")

                st.markdown("**📋 持仓明细**")
                for _stem in _selected_pf_stems:
                    _pf = _b_pf_data.get(_stem) or _pt_load(_stem)
                    if not _pf:
                        continue
                    _dn = _pf.get("portfolio_name", _stem)
                    _sv = _pf.get("saved_at", "")[:10]
                    with st.expander(f"📂 {_dn}  [{_sv}]", expanded=False):
                        _rows = []
                        _total_cost_val = 0.0
                        _total_mkt_val = 0.0
                        _total_gain = 0.0
                        for _h in _pf.get("holdings", []):
                            _sym = _h.get("symbol", "").upper()
                            _sh = float(_h.get("shares", 0))
                            _cost = float(_h.get("avg_cost", 0))
                            _cur = _b_prices.get(_sym, _cost)
                            _mkt = _sh * _cur
                            _cost_val = _sh * _cost
                            _ret = (_cur - _cost) / _cost * 100 if _cost else 0
                            _total_cost_val += _cost_val
                            _total_mkt_val += _mkt
                            _total_gain += (_mkt - _cost_val)
                            _rows.append({
                                "Symbol": _sym, "Shares": _sh,
                                "Avg Cost": f"${_cost:.2f}", "Current": f"${_cur:.2f}",
                                "Market Value": f"${_mkt:,.0f}", "Return %": f"{_ret:+.1f}%",
                                "Note": _h.get("rationale", ""),
                            })
                        _cash = float(_pf.get("cash", 0))
                        _total_mkt_val += _cash
                        _total_cost_val += _cash
                        _total_ret_pct = _total_gain / (_total_cost_val - _cash) * 100 if (_total_cost_val - _cash) > 0 else 0
                        # CASH 行
                        _rows.append({
                            "Symbol": "💵 CASH", "Shares": "—",
                            "Avg Cost": "—", "Current": "—",
                            "Market Value": f"${_cash:,.0f}", "Return %": "0.0%",
                            "Note": "",
                        })
                        # TOTAL 行
                        _rows.append({
                            "Symbol": "📊 TOTAL", "Shares": "—",
                            "Avg Cost": "—", "Current": "—",
                            "Market Value": f"${_total_mkt_val:,.0f}",
                            "Return %": f"{_total_ret_pct:+.1f}%",
                            "Note": "",
                        })
                        if _rows:
                            st.dataframe(pd.DataFrame(_rows), hide_index=True, use_container_width=True)

    # ══════════════════════════════════════════════════════════════
    # MODULE C — 历史回溯归因（多策略对比）
    # ══════════════════════════════════════════════════════════════
    with arena_tab_c:
        st.markdown("### 🔍 历史回溯与持仓归因分析（多策略对比）")
        _all_pf_files_c = _pt_list_files()
        if not _all_pf_files_c:
            st.info("仓库为空，请先在 Tab A 创建策略。")
        else:
            _c_label_to_stem = {f["label"]: f["stem"] for f in _all_pf_files_c}
            _c_labels = list(_c_label_to_stem.keys())

            _c_date_col1, _c_date_col2 = st.columns(2)
            with _c_date_col1:
                _c_start = st.date_input("📅 Start Date", value=pd.Timestamp.now() - pd.Timedelta(days=90), key="bt_start")
            with _c_date_col2:
                _c_end = st.date_input("📅 End Date", value=pd.Timestamp.now(), key="bt_end")

            _c_selected_labels = st.multiselect(
                "🔍 选择策略进行回溯对比（最多 10 个）",
                options=_c_labels,
                default=_c_labels[:min(2, len(_c_labels))],
                max_selections=10,
                key="bt_pf_multiselect",
            )
            _c_selected_stems = [_c_label_to_stem[l] for l in _c_selected_labels]

            if not _c_selected_stems:
                st.info("请选择至少一个策略。")
            elif st.button("📊 生成回溯报告", type="primary", key="bt_run_btn"):
                # 收集所有策略数据
                _c_pf_map: dict[str, dict] = {}
                _c_all_syms: set = set()
                for _stem in _c_selected_stems:
                    _pf = _pt_load(_stem)
                    if _pf:
                        _c_pf_map[_stem] = _pf
                        for _h in _pf.get("holdings", []):
                            if float(_h.get("shares", 0)) > 0:
                                _c_all_syms.add(_h["symbol"].upper())

                if not _c_pf_map:
                    st.error("无法读取策略文件。")
                elif not _c_all_syms:
                    st.warning("所选策略均无持仓。")
                else:
                    _bm_syms = ["XBI", "SPY"]
                    _all_download_syms = list(_c_all_syms) + _bm_syms
                    with st.spinner(f"拉取 {len(_c_all_syms)} 只持仓股 + XBI/SPY 历史数据…"):
                        try:
                            _c_hist_px = yf.download(
                                _all_download_syms, start=str(_c_start), end=str(_c_end),
                                auto_adjust=True, progress=False,
                            )
                            if isinstance(_c_hist_px.columns, pd.MultiIndex):
                                _c_close = _c_hist_px["Close"]
                            else:
                                _c_close = _c_hist_px[["Close"]]
                                _c_close.columns = _all_download_syms

                            # ── 多策略总市值对比曲线 + XBI/SPY 基准 ──
                            _c_portfolio_curves: dict[str, pd.Series] = {}
                            _c_initial_caps: dict[str, float] = {}
                            for _stem, _pf in _c_pf_map.items():
                                _pf_name = _pf.get("portfolio_name", _stem)
                                _syms = [h["symbol"].upper() for h in _pf.get("holdings", []) if float(h.get("shares", 0)) > 0]
                                _shares_map = {h["symbol"].upper(): float(h.get("shares", 0)) for h in _pf.get("holdings", [])}
                                _cash = float(_pf.get("cash", 0))
                                _curve = sum(
                                    _c_close[sym] * _shares_map.get(sym, 0)
                                    for sym in _syms if sym in _c_close.columns
                                )
                                if isinstance(_curve, pd.Series):
                                    _curve = (_curve + _cash).dropna()
                                    _c_portfolio_curves[_pf_name] = _curve
                                    _c_initial_caps[_pf_name] = float(_pf.get("total_capital") or (_curve.iloc[0] if len(_curve) > 0 else 0))

                            if _c_portfolio_curves:
                                _curve_df = pd.DataFrame(_c_portfolio_curves).dropna(how="all")

                                # 为每个策略分别附上同等资金的 XBI/SPY 基准
                                # 若多策略则取平均初始资金做一条共用基准线
                                _avg_cap_c = sum(_c_initial_caps.values()) / max(len(_c_initial_caps), 1)
                                for _bm in _bm_syms:
                                    if _bm in _c_close.columns:
                                        _bm_s = _c_close[_bm].dropna()
                                        if len(_bm_s) > 0 and float(_bm_s.iloc[0]) > 0:
                                            _bm_norm = _bm_s / float(_bm_s.iloc[0]) * _avg_cap_c
                                            _bm_norm.name = f"{_bm} (基准)"
                                            _curve_df = _curve_df.join(_bm_norm, how="left")

                                st.markdown("**📈 多策略总市值变化曲线（含 Cash + XBI/SPY 基准）**")
                                st.line_chart(_curve_df, height=380)

                                # 期间收益率汇总（含 XBI/SPY）
                                _n_metrics = len(_c_portfolio_curves) + len(_bm_syms)
                                _ret_cols = st.columns(min(_n_metrics, 6))
                                _col_idx = 0
                                for _pf_name, _curve in _c_portfolio_curves.items():
                                    if len(_curve) >= 2:
                                        _c_ret = (_curve.iloc[-1] - _curve.iloc[0]) / _curve.iloc[0] * 100
                                        with _ret_cols[_col_idx % 6]:
                                            st.metric(_pf_name, f"${_curve.iloc[-1]:,.0f}", f"{_c_ret:+.2f}%")
                                        _col_idx += 1
                                # XBI / SPY metric
                                for _bm in _bm_syms:
                                    if _bm in _c_close.columns:
                                        _bm_s = _c_close[_bm].dropna()
                                        if len(_bm_s) >= 2:
                                            _bm_ret = (_bm_s.iloc[-1] - _bm_s.iloc[0]) / _bm_s.iloc[0] * 100
                                            _bm_end_val = float(_bm_s.iloc[-1]) / float(_bm_s.iloc[0]) * _avg_cap_c
                                            with _ret_cols[_col_idx % 6]:
                                                st.metric(f"{_bm} (基准)", f"${_bm_end_val:,.0f}", f"{_bm_ret:+.2f}%")
                                            _col_idx += 1

                            # ── 每个策略的个股归因明细（含 CASH + XBI/SPY 对比 + TOTAL 行）──
                            st.markdown("---")
                            st.markdown("**📋 各策略个股归因明细**")
                            for _stem in _c_selected_stems:
                                _pf = _c_pf_map.get(_stem)
                                if not _pf:
                                    continue
                                _pf_name = _pf.get("portfolio_name", _stem)
                                _sv = _pf.get("saved_at", "")[:10]
                                _syms = [h["symbol"].upper() for h in _pf.get("holdings", []) if float(h.get("shares", 0)) > 0]
                                _shares_map = {h["symbol"].upper(): float(h.get("shares", 0)) for h in _pf.get("holdings", [])}
                                _costs_map = {h["symbol"].upper(): float(h.get("avg_cost", 0)) for h in _pf.get("holdings", [])}
                                _cash = float(_pf.get("cash", 0))
                                _cap = float(_pf.get("total_capital") or 0)

                                with st.expander(f"📂 {_pf_name}  [{_sv}]", expanded=len(_c_selected_stems) == 1):
                                    _attr_rows = []
                                    _total_start_val = 0.0
                                    _total_end_val = 0.0
                                    _total_gain = 0.0
                                    for _sym in _syms:
                                        if _sym not in _c_close.columns:
                                            continue
                                        _px_ser = _c_close[_sym].dropna()
                                        if len(_px_ser) < 2:
                                            continue
                                        _p0, _p1 = float(_px_ser.iloc[0]), float(_px_ser.iloc[-1])
                                        _sh = _shares_map.get(_sym, 0)
                                        _start_v = _p0 * _sh
                                        _end_v = _p1 * _sh
                                        _gain = _end_v - _start_v
                                        _ret_pct = (_p1 - _p0) / _p0 * 100 if _p0 else 0
                                        _total_start_val += _start_v
                                        _total_end_val += _end_v
                                        _total_gain += _gain
                                        _attr_rows.append({
                                            "Symbol": _sym,
                                            "Avg Cost": f"${_costs_map.get(_sym, 0):.2f}",
                                            "期初价格": f"${_p0:.2f}",
                                            "期末价格": f"${_p1:.2f}",
                                            "期初市值": f"${_start_v:,.0f}",
                                            "期末市值": f"${_end_v:,.0f}",
                                            "涨跌幅": f"{_ret_pct:+.1f}%",
                                            "盈亏贡献 ($)": f"${_gain:+,.0f}",
                                        })
                                    # CASH 行
                                    _total_start_val += _cash
                                    _total_end_val += _cash
                                    _attr_rows.append({
                                        "Symbol": "💵 CASH",
                                        "Avg Cost": "—",
                                        "期初价格": "—", "期末价格": "—",
                                        "期初市值": f"${_cash:,.0f}",
                                        "期末市值": f"${_cash:,.0f}",
                                        "涨跌幅": "0.0%",
                                        "盈亏贡献 ($)": "$0",
                                    })
                                    # TOTAL 行
                                    _total_ret = _total_gain / _total_start_val * 100 if _total_start_val > 0 else 0
                                    _attr_rows.append({
                                        "Symbol": "📊 TOTAL",
                                        "Avg Cost": "—",
                                        "期初价格": "—", "期末价格": "—",
                                        "期初市值": f"${_total_start_val:,.0f}",
                                        "期末市值": f"${_total_end_val:,.0f}",
                                        "涨跌幅": f"{_total_ret:+.1f}%",
                                        "盈亏贡献 ($)": f"${_total_gain:+,.0f}",
                                    })
                                    # XBI / SPY 基准对比行（同等总资金可比）
                                    _ref_cap = _cap if _cap > 0 else _total_start_val
                                    for _bm in _bm_syms:
                                        if _bm in _c_close.columns:
                                            _bm_s = _c_close[_bm].dropna()
                                            if len(_bm_s) >= 2:
                                                _bp0, _bp1 = float(_bm_s.iloc[0]), float(_bm_s.iloc[-1])
                                                _bm_ret_pct = (_bp1 - _bp0) / _bp0 * 100 if _bp0 else 0
                                                _bm_end_v = _ref_cap * (_bp1 / _bp0) if _bp0 else _ref_cap
                                                _bm_gain = _bm_end_v - _ref_cap
                                                _attr_rows.append({
                                                    "Symbol": f"📌 {_bm} 基准",
                                                    "Avg Cost": "—",
                                                    "期初价格": f"${_bp0:.2f}",
                                                    "期末价格": f"${_bp1:.2f}",
                                                    "期初市值": f"${_ref_cap:,.0f}",
                                                    "期末市值": f"${_bm_end_v:,.0f}",
                                                    "涨跌幅": f"{_bm_ret_pct:+.1f}%",
                                                    "盈亏贡献 ($)": f"${_bm_gain:+,.0f}",
                                                })
                                    if _attr_rows:
                                        st.dataframe(
                                            pd.DataFrame(_attr_rows),
                                            hide_index=True, use_container_width=True,
                                        )

                        except Exception as _bt_err:
                            st.error(f"回溯数据拉取失败：{_bt_err}")

# ── Footer ────────────────────────────────────────────────────────

st.markdown("<div class='footer'>Dashboard by Tony Jiang</div>", unsafe_allow_html=True)
